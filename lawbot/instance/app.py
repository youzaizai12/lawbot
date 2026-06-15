from flask import Flask, render_template, request, jsonify, send_file
from flask_sqlalchemy import SQLAlchemy
from datetime import datetime, date
import requests
import json
import time
import threading
from collections import defaultdict
import io
import re
import os

# ====================== 导入本地知识库模块 ======================
from knowledge_base.knowledge_manager import KnowledgeManager

# ====================== 初始化本地知识库 ======================
print("正在加载本地知识库...")
knowledge_manager = KnowledgeManager()
print("本地知识库加载完成")

app = Flask(__name__)

# 数据库配置
app.config['SQLALCHEMY_DATABASE_URI'] = 'sqlite:///chat.db'
app.config['SQLALCHEMY_TRACK_MODIFICATIONS'] = False
db = SQLAlchemy(app)
from flask_migrate import Migrate

migrate = Migrate(app, db)


# 对话日志表
class ChatLog(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    username = db.Column(db.String(50), default="user")
    user_msg = db.Column(db.Text, nullable=False)
    bot_msg = db.Column(db.Text, nullable=False)
    context = db.Column(db.Text, default="[]")
    scene = db.Column(db.String(20), default="auto", index=True)
    is_pro = db.Column(db.Boolean, default=False)
    risk_level = db.Column(db.String(20), default="low")
    risk_alert = db.Column(db.Text, default="")
    rag_sources = db.Column(db.Text, default="")
    retrieval_count = db.Column(db.Integer, default=0)
    create_time = db.Column(db.DateTime, default=datetime.now, index=True)
    __table_args__ = (
        db.Index('idx_risk_create', 'risk_level', 'create_time'),
        db.Index('idx_pro_scene', 'is_pro', 'scene'),
    )


# 压力测试结果存储
test_results = {
    "total_requests": 0,
    "success_requests": 0,
    "fail_requests": 0,
    "avg_response_time": 0.0,
    "response_times": [],
    "error_details": defaultdict(int)
}

# ====================== 14类法律场景关键词配置 ======================

# 1. 民间借贷
PRIVATE_LENDING_KEYWORDS = ["借钱", "欠钱", "借条", "欠条", "转账", "债务", "还款", "利息", "高利贷",
                            "担保", "保证人", "催收", "讨债", "借贷", "出借", "借款人", "贷款"]

# 2. 劳动纠纷
LABOR_KEYWORDS = ["工资", "辞退", "拖欠", "加班", "社保", "工伤", "劳动合同", "仲裁", "工地",
                  "赔偿金", "经济补偿", "N+1", "2N", "违法解除", "竞业限制", "劳务派遣", "试用期",
                  "开除", "离职", "调岗", "降薪", "五险一金"]

# 3. 交通事故
TRAFFIC_KEYWORDS = ["车祸", "撞车", "肇事", "逃逸", "酒驾", "醉驾", "交强险", "三者险", "理赔",
                    "伤残鉴定", "误工费", "医药费", "交通肇事", "开车", "电动车", "摩托车", "行人"]

# 4. 婚姻家庭
MARRIAGE_KEYWORDS = ["离婚", "抚养权", "彩礼", "结婚", "财产分割", "家暴", "家庭暴力", "感情破裂",
                     "重婚", "同居", "婚外情", "出轨", "抚养费", "赡养", "探望权", "婚内协议"]

# 5. 刑事法律
CRIMINAL_KEYWORDS = ["盗窃", "诈骗", "轻伤", "拘留", "逮捕", "取保", "缓刑", "坐牢", "立案",
                     "罪名", "故意伤害", "抢劫", "强奸", "贪污", "受贿", "帮信", "掩饰隐瞒",
                     "取保候审", "监视居住", "刑事", "判刑", "监狱", "刑事责任"]

# 6. 房产纠纷
PROPERTY_KEYWORDS = ["买房", "卖房", "租房", "房东", "租客", "押金", "房产证", "产权", "物业",
                     "开发商", "预售", "定金", "首付", "贷款买房", "二手房", "中介", "过户",
                     "房屋质量", "漏水", "违建", "拆迁"]

# 7. 合同纠纷
CONTRACT_KEYWORDS = ["合同", "违约", "解除合同", "合同无效", "格式条款", "定金", "订金", "霸王条款",
                     "合同纠纷", "协议", "承诺", "要约", "买卖合同", "服务合同", "租赁合同"]

# 8. 侵权责任
TORT_KEYWORDS = ["侵权", "赔偿", "人身损害", "名誉权", "隐私权", "肖像权", "姓名权", "高空抛物",
                 "动物伤人", "狗咬", "环境污染", "网络侵权", "诽谤", "侮辱"]

# 9. 诉讼程序
PROCEDURE_KEYWORDS = ["起诉", "上诉", "诉讼", "法院", "开庭", "判决", "执行", "仲裁", "诉讼时效",
                      "管辖", "立案", "证据", "举证", "证人", "诉讼费", "律师费", "再审", "申诉"]

# 10. 公司法律
CORPORATE_KEYWORDS = ["公司", "股东", "股权", "法人", "法定代表人", "破产", "清算", "注销",
                      "增资", "减资", "章程", "董事会", "监事会", "合伙", "合伙企业"]

# 11. 知识产权
IP_KEYWORDS = ["专利", "商标", "版权", "著作权", "侵权", "盗版", "山寨", "知识产权",
               "商业秘密", "技术秘密", "注册", "申请专利", "注册商标"]

# 12. 消费者权益
CONSUMER_KEYWORDS = ["假货", "退货", "退款", "欺诈", "三倍赔偿", "十倍赔偿", "消费者",
                     "维权", "12315", "预付卡", "会员卡", "霸王条款", "食品安全"]

# 13. 行政法律
ADMINISTRATIVE_KEYWORDS = ["行政拘留", "行政处罚", "行政复议", "行政诉讼", "政府信息",
                           "公开", "拆迁补偿", "行政许可", "行政强制", "罚款", "吊销执照"]

# 14. 继承法律
INHERITANCE_KEYWORDS = ["继承", "遗嘱", "遗产", "法定继承", "遗赠", "公证", "继承人",
                        "放弃继承", "遗产分割", "财产分配", "死亡"]

# 场景映射表
SCENE_CONFIG = {
    "private_lending": {
        "name": "民间借贷",
        "keywords": PRIVATE_LENDING_KEYWORDS,
        "color": "#e8f5e9",
        "text_color": "#2e7d32"
    },
    "labor": {
        "name": "劳动纠纷",
        "keywords": LABOR_KEYWORDS,
        "color": "#fff3e0",
        "text_color": "#e65100"
    },
    "traffic_accident": {
        "name": "交通事故",
        "keywords": TRAFFIC_KEYWORDS,
        "color": "#e3f2fd",
        "text_color": "#1565c0"
    },
    "marriage_family": {
        "name": "婚姻家庭",
        "keywords": MARRIAGE_KEYWORDS,
        "color": "#fce4ec",
        "text_color": "#c2185b"
    },
    "criminal": {
        "name": "刑事法律",
        "keywords": CRIMINAL_KEYWORDS,
        "color": "#ffebee",
        "text_color": "#c62828"
    },
    "property": {
        "name": "房产纠纷",
        "keywords": PROPERTY_KEYWORDS,
        "color": "#e8eaf6",
        "text_color": "#283593"
    },
    "contract": {
        "name": "合同纠纷",
        "keywords": CONTRACT_KEYWORDS,
        "color": "#f1f8e9",
        "text_color": "#33691e"
    },
    "tort": {
        "name": "侵权责任",
        "keywords": TORT_KEYWORDS,
        "color": "#f3e5f5",
        "text_color": "#6a1b9a"
    },
    "procedure": {
        "name": "诉讼程序",
        "keywords": PROCEDURE_KEYWORDS,
        "color": "#e0f7fa",
        "text_color": "#006064"
    },
    "corporate": {
        "name": "公司法律",
        "keywords": CORPORATE_KEYWORDS,
        "color": "#efebe9",
        "text_color": "#4e342e"
    },
    "intellectual_property": {
        "name": "知识产权",
        "keywords": IP_KEYWORDS,
        "color": "#e8eaf6",
        "text_color": "#1a237e"
    },
    "consumer": {
        "name": "消费者权益",
        "keywords": CONSUMER_KEYWORDS,
        "color": "#fff8e1",
        "text_color": "#f57f17"
    },
    "administrative": {
        "name": "行政法律",
        "keywords": ADMINISTRATIVE_KEYWORDS,
        "color": "#e0f2f1",
        "text_color": "#004d40"
    },
    "inheritance": {
        "name": "继承法律",
        "keywords": INHERITANCE_KEYWORDS,
        "color": "#fbe9e7",
        "text_color": "#bf360c"
    }
}

# 场景名称映射（用于前端显示）
SCENE_NAMES = {key: config["name"] for key, config in SCENE_CONFIG.items()}

# ====================== 提示词模板 ======================
BASE_SYSTEM_PROMPT = """
你是专业的法律咨询助手。

【回答要求】
1. 优先使用下方知识库中检索到的信息回答问题
2. 法律分析仅作参考，必须提示"不可替代律师专业意见"
3. 回答控制在200字左右，简洁清晰
4. 回答完后推荐3个相关问题

格式：回答内容 + 换行 + 💡你可以继续问：问题1？问题2？问题3？
"""

PRO_SYSTEM_PROMPT = """
你是资深法律专家，为同行提供专业支持。

【重要】你只需要回答法律内容，不要在回答末尾添加类似"📚 知识库检索"、"📚 【知识库检索】"等来源标记，这些会由系统自动添加。

【回答格式】
**一、法律依据与构成要件**
**二、举证责任分配**
**三、诉讼策略与风险**
**四、类案裁判观点**
**五、实操建议**
**六、专业延伸（3个问题）**
"""


def clean_rag_markers(text):
    """清理文本中可能自带的RAG来源标记，避免重复显示"""
    if not text:
        return text

    # 移除各种格式的RAG标记
    patterns = [
        r'\n*---\n*📚\s*【知识库检索】[^\n]*\n*',
        r'\n*---\n*📚\s*知识库检索[^\n]*\n*',
        r'\n*---\n*【知识库检索】[^\n]*\n*',
        r'\n*---\n*知识库检索[^\n]*\n*',
        r'\n*📚\s*【知识库检索】[^\n]*\n*',
        r'\n*📚\s*知识库检索[^\n]*\n*',
    ]

    for pattern in patterns:
        text = re.sub(pattern, '', text, flags=re.DOTALL)

    # 清理多余的空行
    text = re.sub(r'\n{3,}', '\n\n', text)

    return text.strip()


def get_rag_context(user_msg: str):
    """从本地知识库检索相关内容"""
    try:
        results = knowledge_manager.search(user_msg, top_k=3)

        if not results:
            return None, []

        context_parts = []
        sources = []

        for result in results:
            category = result.get("metadata", {}).get("category", "未知")
            content = result.get("content", "")
            sources.append(category)

            if len(content) > 500:
                content = content[:500] + "..."

            context_parts.append(f"【{category}法律知识】\n{content}")

        rag_text = "\n\n---\n\n".join(context_parts)
        return rag_text, sources

    except Exception as e:
        print(f"RAG检索失败: {e}")
        return None, []


# ====================== 14类场景识别函数 ======================
def auto_detect_scene(user_msg: str) -> str:
    """自动识别14类法律场景"""
    msg_lower = user_msg.lower()

    # 计算每个场景的匹配分数
    scene_scores = {}
    for scene_key, config in SCENE_CONFIG.items():
        score = sum(1 for kw in config["keywords"] if kw in msg_lower)
        scene_scores[scene_key] = score

    # 找出最高分
    max_score = max(scene_scores.values()) if scene_scores else 0

    # 如果最高分 >= 2，返回匹配的场景
    if max_score >= 2:
        for scene_key, score in scene_scores.items():
            if score == max_score:
                return scene_key

    # 如果最高分 = 1，检查是否有明确的单一场景
    if max_score == 1:
        matched_scenes = [k for k, v in scene_scores.items() if v == 1]
        if len(matched_scenes) == 1:
            return matched_scenes[0]

    # 默认返回合同纠纷（最常见的民事场景）
    return "contract"


def get_scene_name(scene_key: str) -> str:
    """获取场景的中文名称"""
    return SCENE_NAMES.get(scene_key, "法律咨询")


def get_scene_css_class(scene_key: str) -> str:
    """获取场景对应的CSS类名"""
    return f"scene-tag {scene_key}"


# ====================== 混合风险评估 ======================
def detect_risk_level(user_msg: str, scene: str) -> tuple:
    """混合风险评估：关键词 + 知识库 + 场景辅助"""
    msg_lower = user_msg.lower()

    # ========== 第一层：关键词快速匹配（高风险） ==========
    high_keywords = [
        "家暴", "家庭暴力", "拘留", "刑事拘留", "逮捕", "判刑", "坐牢",
        "死亡", "重伤", "杀人", "打死", "生命危险", "逃跑", "紧急",
        "百万", "千万", "巨额", "全部积蓄", "养老钱", "救命钱"
    ]
    for kw in high_keywords:
        if kw in msg_lower:
            return "high", "⚠️ 【高风险】涉及重大法律风险，建议立即咨询律师"

    # ========== 第二层：知识库检索验证 ==========
    try:
        results = knowledge_manager.search(user_msg, top_k=3)
        high_count = 0
        medium_count = 0

        for result in results:
            category = result.get("metadata", {}).get("category", "")
            content = result.get("content", "").lower()

            # 根据分类和内容判断风险等级
            if category in ["刑事法律", "刑事"]:
                high_count += 1
            elif category in ["交通事故", "交通"]:
                if "死亡" in content or "逃逸" in content or "重伤" in content:
                    high_count += 1
                else:
                    medium_count += 1
            elif category in ["劳动纠纷", "劳动"]:
                if "工伤" in content or "死亡" in content or "伤残" in content:
                    high_count += 1
                elif "拖欠" in content or "辞退" in content:
                    medium_count += 1
            elif category in ["民间借贷", "婚姻家庭", "合同纠纷", "房产纠纷", "婚姻", "合同", "房产"]:
                if "诉讼" in content or "法院" in content or "起诉" in content:
                    medium_count += 1
            elif category in ["侵权责任", "诉讼程序"]:
                medium_count += 1

        if high_count >= 1:
            return "high", "⚠️ 【高风险】根据相关案例，您的情况存在较高法律风险，建议立即咨询律师"
        if medium_count >= 2:
            return "medium", "🔔 【中风险】根据类似案例，建议尽快收集证据并咨询专业律师"

    except Exception as e:
        print(f"知识库风险评估失败: {e}")

    # ========== 第三层：场景辅助判断 ==========
    if scene == "criminal":
        return "medium", "🔔 【中风险】涉及刑事法律问题，建议咨询专业律师"

    # 检查中风险关键词
    medium_keywords = [
        "起诉", "诉讼", "法院", "开庭", "判决", "执行", "仲裁",
        "赔偿", "辞退", "拖欠工资", "离婚", "抚养权", "财产分割"
    ]
    for kw in medium_keywords:
        if kw in msg_lower:
            return "medium", "🔔 【中风险】建议尽快收集证据并咨询专业律师"

    return "low", "ℹ️ 【低风险】一般法律咨询，以上分析仅供参考"


def merge_risk_with_reply(bot_reply: str, risk_level: str, risk_alert: str) -> str:
    """合并风险提示"""
    if risk_level == "high":
        return f"{risk_alert}\n\n---\n\n{bot_reply}"
    elif risk_level == "medium":
        return f"{bot_reply}\n\n---\n\n{risk_alert}"
    else:
        return f"{bot_reply}\n\n{risk_alert}"


# ====================== DeepSeek 配置 ======================
DEEPSEEK_API_KEY = "sk-efb0b7c994c84df6aab93c6a66a1ad1f"
DEEPSEEK_URL = "https://api.deepseek.com/v1/chat/completions"

import urllib3

urllib3.disable_warnings(urllib3.exceptions.InsecureRequestWarning)
session = requests.Session()
session.verify = False


def call_deepseek(messages, scene="civil", is_pro=False, retry_times=2, rag_text=None):
    """调用DeepSeek API"""
    scene_name = get_scene_name(scene)

    if is_pro:
        system_prompt = PRO_SYSTEM_PROMPT
    else:
        system_prompt = BASE_SYSTEM_PROMPT

    if rag_text:
        system_prompt = f"""
{system_prompt}

【当前咨询场景】{scene_name}

【知识库检索结果 - 请优先使用以下信息回答用户问题】
{rag_text}

请基于以上知识库信息回答用户问题，确保回答准确、专业。
"""

    messages[0]["content"] = system_prompt
    max_context = 12 if is_pro else 8
    messages = messages[:max_context * 2 + 1]

    headers = {
        "Authorization": f"Bearer {DEEPSEEK_API_KEY}",
        "Content-Type": "application/json"
    }

    payload = {
        "model": "deepseek-chat",
        "messages": messages,
        "temperature": 0.15 if is_pro else 0.25,
        "stream": False
    }

    for retry in range(retry_times + 1):
        try:
            start_time = time.time()
            resp = session.post(DEEPSEEK_URL, headers=headers, json=payload, timeout=35)
            end_time = time.time()

            if resp.status_code == 200:
                data = resp.json()
                return {
                    "reply": data["choices"][0]["message"]["content"],
                    "success": True,
                    "response_time": end_time - start_time
                }
            else:
                if retry < retry_times:
                    time.sleep(0.8 * (retry + 1))
                    continue
                raise Exception(f"API错误 {resp.status_code}")

        except Exception as e:
            if retry < retry_times:
                time.sleep(0.8 * (retry + 1))
                continue
            return {
                "reply": "服务暂时不可用，请稍后重试。",
                "success": False,
                "response_time": 0,
                "error": str(e)
            }

    return {"reply": "服务暂时不可用", "success": False, "response_time": 0}


# ====================== 固定文书模板（完全静态，无占位符） ======================
DOC_TEMPLATES = {
    "complaint": """民事起诉状

原告：待补充（姓名、性别、出生年月、民族、职业、住址、身份证号）

被告：待补充（姓名、性别、出生年月、民族、职业、住址、身份证号）

诉讼请求：
1. 待补充

事实与理由：
待补充

此致
待补充人民法院

具状人：待补充
年  月  日
""",

    "defense": """民事答辩状

答辩人：待补充（姓名、性别、出生年月、住址、身份证号）

针对原告的起诉，答辩如下：
待补充

此致
待补充人民法院

答辩人：待补充
年  月  日
""",

    "labor_application": """劳动仲裁申请书

申请人：待补充（姓名、性别、出生年月、住址、身份证号）

被申请人：待补充（名称、住所地、法定代表人）

仲裁请求：
1. 待补充

事实与理由：
待补充

此致
待补充劳动人事争议仲裁委员会

申请人：待补充
年  月  日
""",

    "contract_clause": """合同条款建议

根据您的需求，建议包含以下核心条款：

1. 待补充
2. 待补充
3. 待补充

注：以上条款仅供参考，实际签订合同时建议咨询专业律师。
""",

    "criminal_defense": """刑事辩护意见书

当事人：待补充
涉嫌罪名：待补充

案件事实概述：
待补充

辩护意见：
1. 待补充
2. 待补充

结论：
待补充

辩护人：待补充
年  月  日
"""
}


# ====================== 生成法律文书（完全静态，不调用AI） ======================
@app.route('/generate-document', methods=['POST'])
def generate_document():
    data = request.get_json()
    doc_type = data.get("doc_type", "complaint")
    doc_category = data.get("doc_category", "")  # 仅用于前端记录，不影响生成

    # 获取固定静态模板
    template = DOC_TEMPLATES.get(doc_type, DOC_TEMPLATES["complaint"])

    # 生成文件名
    doc_names = {
        "complaint": "民事起诉状",
        "defense": "民事答辩状",
        "labor_application": "劳动仲裁申请书",
        "contract_clause": "合同条款建议",
        "criminal_defense": "刑事辩护意见书"
    }
    filename = f"{doc_names.get(doc_type, '法律文书')}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.txt"

    return jsonify({
        "success": True,
        "document": template,
        "filename": filename,
        "rag_used": False,
        "rag_sources": []
    })


@app.route('/export-document', methods=['POST'])
def export_document():
    data = request.get_json()
    content = data.get("content", "")
    filename = data.get("filename", "法律文书.txt")
    file_format = data.get("format", "txt")

    if not content:
        return jsonify({"error": "内容为空"}), 400

    try:
        if file_format == "docx":
            from docx import Document
            from docx.shared import Pt, Inches
            from docx.enum.text import WD_ALIGN_PARAGRAPH

            doc = Document()
            section = doc.sections[0]
            section.top_margin = Inches(1)
            section.bottom_margin = Inches(1)
            section.left_margin = Inches(1.2)
            section.right_margin = Inches(1.2)

            title = doc.add_heading(filename.replace('.docx', ''), level=1)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER

            for line in content.split('\n'):
                if line.strip():
                    p = doc.add_paragraph(line.strip())
                    if p.runs:
                        p.runs[0].font.size = Pt(11)

            file_stream = io.BytesIO()
            doc.save(file_stream)
            file_stream.seek(0)

            return send_file(
                file_stream,
                as_attachment=True,
                download_name=filename,
                mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
            )
        else:
            file_stream = io.BytesIO(content.encode('utf-8'))
            return send_file(
                file_stream,
                as_attachment=True,
                download_name=filename.replace('.docx', '.txt'),
                mimetype='text/plain; charset=utf-8'
            )
    except Exception as e:
        return jsonify({"error": f"导出失败: {str(e)}"}), 500


# ====================== 路由 ======================
@app.route('/')
def index():
    return render_template('index.html')


# ====================== 聊天路由 ======================
@app.route('/chat', methods=['POST'])
def chat():
    data = request.get_json()
    user_msg = data.get("msg", "").strip()
    context = data.get("context", [])

    if not user_msg:
        return jsonify({"reply": "请输入法律问题", "context": [], "success": False})

    # 1. 从知识库检索相关内容
    rag_text, rag_sources = get_rag_context(user_msg)

    # 2. 如果知识库检索不到相关内容，拒绝回答
    if not rag_sources:
        return jsonify({
            "reply": "抱歉，我仅能回答法律相关问题。您问的问题不在我的知识库范围内，请提出具体的法律咨询问题。\n\n💡 您可以问：民间借贷、劳动纠纷、交通事故、婚姻家庭、刑事法律等问题。",
            "context": context,
            "success": False
        })

    # 3. 场景识别（14类）
    scene = auto_detect_scene(user_msg)

    # 4. 混合风险评估
    risk_level, risk_alert = detect_risk_level(user_msg, scene)

    # 5. 构建对话消息
    messages = [{"role": "system", "content": ""}]
    for item in context[-10:]:
        messages.append({"role": "user", "content": item["user"]})
        messages.append({"role": "assistant", "content": item["bot"]})
    messages.append({"role": "user", "content": user_msg})

    # 6. 调用AI
    result = call_deepseek(messages, scene, is_pro=False, rag_text=rag_text)

    # 7. 清理模型输出中可能自带的RAG标记
    bot_reply_raw = clean_rag_markers(result["reply"])

    # 8. 合并风险提示
    bot_reply = merge_risk_with_reply(bot_reply_raw, risk_level, risk_alert)

    # 9. 添加知识库来源标记（只加一次）
    if rag_sources:
        rag_footer = f"\n\n---\n📚 【知识库检索】参考了 {len(rag_sources)} 个知识条目：{', '.join(rag_sources)}"
        bot_reply += rag_footer

    # 10. 保存对话
    new_context = context[-10:]
    new_context.append({"user": user_msg, "bot": bot_reply})

    try:
        log = ChatLog(
            user_msg=user_msg,
            bot_msg=bot_reply,
            context=json.dumps(new_context, ensure_ascii=False),
            scene=scene,
            is_pro=False,
            risk_level=risk_level,
            risk_alert=risk_alert[:300],
            rag_sources=",".join(rag_sources) if rag_sources else "",
            retrieval_count=len(rag_sources) if rag_sources else 0
        )
        db.session.add(log)
        db.session.commit()
    except Exception as e:
        db.session.rollback()
        print(f"数据库保存失败: {e}")

    return jsonify({
        "reply": bot_reply,
        "context": new_context,
        "success": result["success"],
        "rag_used": bool(rag_sources),
        "rag_sources": rag_sources,
        "scene": scene,
        "scene_name": get_scene_name(scene)
    })


# 专业模式聊天
@app.route('/chat-pro', methods=['POST'])
def chat_pro():
    data = request.get_json()
    user_msg = data.get("msg", "").strip()
    context = data.get("context", [])

    if not user_msg:
        return jsonify({"reply": "请输入专业法律问题", "context": [], "success": False})

    # 1. 从知识库检索相关内容
    rag_text, rag_sources = get_rag_context(user_msg)

    # 2. 如果知识库检索不到相关内容，拒绝回答
    if not rag_sources:
        return jsonify({
            "reply": "抱歉，我仅能回答法律相关问题。您问的问题不在我的知识库范围内，请提出具体的法律咨询问题。",
            "context": context,
            "success": False
        })

    # 3. 场景识别（14类）
    scene = auto_detect_scene(user_msg)

    # 4. 混合风险评估
    risk_level, risk_alert = detect_risk_level(user_msg, scene)

    if risk_level == "high":
        risk_alert = "⚠️ 高风险案件，建议立即启动应急法律程序。"
    elif risk_level == "medium":
        risk_alert = "🔔 中风险案件，建议尽快安排法律行动。"
    else:
        risk_alert = "ℹ️ 一般咨询，标准处理流程。"

    # 5. 构建对话消息
    messages = [{"role": "system", "content": ""}]
    for item in context[-12:]:
        messages.append({"role": "user", "content": item["user"]})
        messages.append({"role": "assistant", "content": item["bot"]})
    messages.append({"role": "user", "content": user_msg})

    # 6. 调用AI
    result = call_deepseek(messages, scene, is_pro=True, rag_text=rag_text)

    # 7. 清理模型输出中可能自带的RAG标记
    bot_reply_raw = clean_rag_markers(result["reply"])

    # 8. 合并风险提示
    bot_reply = merge_risk_with_reply(bot_reply_raw, risk_level, risk_alert)

    # 9. 添加知识库来源标记（只加一次）
    if rag_sources:
        rag_footer = f"\n\n---\n📚 【知识库检索】专业模式参考了 {len(rag_sources)} 个知识条目：{', '.join(rag_sources)}"
        bot_reply += rag_footer

    # 10. 保存对话
    new_context = context[-12:]
    new_context.append({"user": user_msg, "bot": bot_reply})

    try:
        log = ChatLog(
            user_msg=user_msg,
            bot_msg=bot_reply,
            context=json.dumps(new_context, ensure_ascii=False),
            scene=scene,
            is_pro=True,
            risk_level=risk_level,
            risk_alert=risk_alert[:300],
            rag_sources=",".join(rag_sources) if rag_sources else "",
            retrieval_count=len(rag_sources) if rag_sources else 0
        )
        db.session.add(log)
        db.session.commit()
    except Exception as e:
        db.session.rollback()
        print(f"数据库保存失败: {e}")

    return jsonify({
        "reply": bot_reply,
        "context": new_context,
        "success": result["success"],
        "rag_used": bool(rag_sources),
        "rag_sources": rag_sources,
        "scene": scene,
        "scene_name": get_scene_name(scene)
    })


# ====================== RAG API路由 ======================
@app.route('/rag-search', methods=['POST'])
def rag_search():
    data = request.get_json()
    query = data.get("query", "")

    if not query:
        return jsonify({"error": "查询内容为空"}), 400

    results = knowledge_manager.search(query, top_k=5)

    return jsonify({
        "success": True,
        "query": query,
        "results": results,
        "total_documents": knowledge_manager.vector_store.get_stats()['total_documents']
    })


@app.route('/rag-stats', methods=['GET'])
def rag_stats():
    stats = knowledge_manager.get_stats()
    return jsonify({
        "total_categories": stats['total_categories'],
        "total_cases": stats['total_cases'],
        "total_documents": stats['vector_stats']['total_documents'],
        "vocabulary_size": stats['vector_stats']['vocabulary_size']
    })


@app.route('/system-check')
def system_check():
    stats = knowledge_manager.get_stats()
    return jsonify({
        "rag_initialized": True,
        "knowledge_base_categories": stats['total_categories'],
        "knowledge_base_cases": stats['total_cases'],
        "knowledge_ready": True
    })


# ====================== 统计接口 ======================
@app.route('/pro-stat')
def pro_stat():
    today = date.today()
    total = ChatLog.query.filter(db.func.date(ChatLog.create_time) == today).count()
    pro = ChatLog.query.filter(db.func.date(ChatLog.create_time) == today, ChatLog.is_pro == True).count()
    rag_used = ChatLog.query.filter(db.func.date(ChatLog.create_time) == today, ChatLog.retrieval_count > 0).count()

    civil = ChatLog.query.filter(db.func.date(ChatLog.create_time) == today, ChatLog.scene == "contract").count()
    labor = ChatLog.query.filter(db.func.date(ChatLog.create_time) == today, ChatLog.scene == "labor").count()
    criminal = ChatLog.query.filter(db.func.date(ChatLog.create_time) == today, ChatLog.scene == "criminal").count()

    high_risk = ChatLog.query.filter(db.func.date(ChatLog.create_time) == today, ChatLog.risk_level == "high").count()
    medium_risk = ChatLog.query.filter(db.func.date(ChatLog.create_time) == today,
                                       ChatLog.risk_level == "medium").count()

    stats = knowledge_manager.get_stats()
    all_cnt = civil + labor + criminal
    civil_pct = round(civil / all_cnt * 100, 1) if all_cnt else 0
    labor_pct = round(labor / all_cnt * 100, 1) if all_cnt else 0
    criminal_pct = round(criminal / all_cnt * 100, 1) if all_cnt else 0

    return jsonify({
        "today_total": total,
        "today_pro": pro,
        "civil_pct": civil_pct,
        "labor_pct": labor_pct,
        "criminal_pct": criminal_pct,
        "high_risk": high_risk,
        "medium_risk": medium_risk,
        "rag_used_today": rag_used,
        "knowledge_base_categories": stats['total_categories'],
        "knowledge_base_cases": stats['total_cases']
    })


@app.route('/risk-logs')
def risk_logs():
    logs = ChatLog.query.filter(ChatLog.risk_level.in_(['high', 'medium'])).order_by(ChatLog.create_time.desc()).limit(
        50).all()
    result = []
    for log in logs:
        result.append({
            "id": log.id,
            "user_msg": log.user_msg[:200],
            "risk_level": log.risk_level,
            "risk_alert": log.risk_alert[:300],
            "scene": log.scene,
            "create_time": log.create_time.strftime("%Y-%m-%d %H:%M:%S"),
            "rag_sources": log.rag_sources
        })
    return jsonify(result)


# ====================== 压力测试 ======================
@app.route('/stress-test', methods=['POST'])
def stress_test():
    global test_results
    test_results = {
        "total_requests": 0,
        "success_requests": 0,
        "fail_requests": 0,
        "avg_response_time": 0.0,
        "response_times": [],
        "error_details": defaultdict(int)
    }

    data = request.get_json()
    concurrency = int(data.get("concurrency", 10))
    total = int(data.get("total", 10))
    test_msg = data.get("test_msg", "劳动合同纠纷如何维权")

    def test_single_request():
        global test_results
        start_time = time.time()
        messages = [{"role": "system", "content": "你是法律助手"}]
        messages.append({"role": "user", "content": test_msg})
        rag_text, _ = get_rag_context(test_msg)
        result = call_deepseek(messages, "labor", is_pro=False, rag_text=rag_text)
        end_time = time.time()

        test_results["total_requests"] += 1
        if result["success"]:
            test_results["success_requests"] += 1
            test_results["response_times"].append(end_time - start_time)
        else:
            test_results["fail_requests"] += 1

        if test_results["response_times"]:
            test_results["avg_response_time"] = sum(test_results["response_times"]) / len(
                test_results["response_times"])

    threads = []
    for i in range(total):
        while len(threading.enumerate()) > concurrency:
            time.sleep(0.01)
        t = threading.Thread(target=test_single_request)
        threads.append(t)
        t.start()

    for t in threads:
        t.join()

    return jsonify(test_results)


@app.route('/test-result')
def get_test_result():
    return jsonify(test_results)


@app.route('/clear-test-result', methods=['POST'])
def clear_test_result():
    global test_results
    test_results = {
        "total_requests": 0,
        "success_requests": 0,
        "fail_requests": 0,
        "avg_response_time": 0.0,
        "response_times": [],
        "error_details": defaultdict(int)
    }
    return jsonify({"status": "cleared"})


# 初始化数据库
with app.app_context():
    db.create_all()

if __name__ == '__main__':
    print("=" * 70)
    print("🤖 智能法律咨询机器人 - 本地知识库RAG系统")
    print("=" * 70)
    print(f"📚 支持14类法律场景识别：")
    for key, config in SCENE_CONFIG.items():
        print(f"   - {config['name']} ({len(config['keywords'])}个关键词)")
    print("=" * 70)

    stats = knowledge_manager.get_stats()
    print(f"📚 知识库状态: 已加载")
    print(f"📂 知识分类数: {stats['total_categories']}")
    print(f"📄 案例总数: {stats['total_cases']}")
    print(f"🔢 向量片段数: {stats['vector_stats']['total_documents']}")
    print(f"📊 词汇表大小: {stats['vector_stats']['vocabulary_size']}")
    print(f"🔍 RAG模式: 本地知识库 → 向量检索 → AI增强回答")
    print(f"⚠️ 风险评估: 混合模式（关键词 + 知识库 + 场景）")
    print(f"📝 文书生成: 完全静态模板（不调用AI）")
    print("=" * 70)
    print("✅ 系统启动成功！访问 http://localhost:5004")
    print("=" * 70)
    app.run(debug=True, host='0.0.0.0', port=5004)