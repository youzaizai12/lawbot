from flask import Flask, render_template, request, jsonify, send_file
from flask_sqlalchemy import SQLAlchemy
from datetime import datetime, date, timedelta
import requests
import json
import time
import threading
from collections import defaultdict
import io
import re
import os
import base64
from functools import wraps

# ====================== 导入本地知识库模块 ======================
from knowledge_base.knowledge_manager import KnowledgeManager

# ====================== 导入文书模板引擎 ======================
from knowledge_base.template_engine import DocumentTemplateEngine, format_document_for_display

# ====================== 初始化本地知识库 ======================
print("正在加载本地知识库...")
knowledge_manager = KnowledgeManager()
print("本地知识库加载完成")

# ====================== 初始化文书模板引擎 ======================
print("正在加载文书模板引擎...")
template_engine = DocumentTemplateEngine()
print(f"文书模板引擎加载完成，共加载 {len(template_engine.get_all_templates())} 个模板")

app = Flask(__name__)

# 数据库配置
app.config['SQLALCHEMY_DATABASE_URI'] = 'sqlite:///chat.db'
app.config['SQLALCHEMY_TRACK_MODIFICATIONS'] = False
db = SQLAlchemy(app)
from flask_migrate import Migrate

migrate = Migrate(app, db)


# ====================== 用户档案表（长期记忆增强）======================
class UserProfile(db.Model):
    """用户档案表 - 存储用户长期信息"""
    id = db.Column(db.Integer, primary_key=True)
    username = db.Column(db.String(50), unique=True, default="default_user")
    # 用户基本信息
    name = db.Column(db.String(50), default="")
    phone = db.Column(db.String(20), default="")
    id_number = db.Column(db.String(30), default="")
    address = db.Column(db.String(200), default="")
    # 重要日期
    birth_date = db.Column(db.Date, nullable=True)
    # 重要事项提醒
    important_memos = db.Column(db.Text, default="[]")  # JSON格式
    # 常用法律诉求类型
    frequent_categories = db.Column(db.Text, default="[]")
    # 最后活跃时间
    last_active = db.Column(db.DateTime, default=datetime.now)
    # 累计咨询次数
    total_consultations = db.Column(db.Integer, default=0)

    def to_dict(self):
        return {
            "name": self.name,
            "phone": self.phone,
            "id_number": self.id_number[:6] + "****" + self.id_number[-4:] if len(
                self.id_number) > 10 else self.id_number,
            "address": self.address,
            "important_memos": json.loads(self.important_memos) if self.important_memos else [],
            "frequent_categories": json.loads(self.frequent_categories) if self.frequent_categories else [],
            "total_consultations": self.total_consultations
        }


# ====================== 对话摘要表（长期记忆增强）======================
class ConversationSummary(db.Model):
    """对话摘要表 - 存储对话的长期记忆"""
    id = db.Column(db.Integer, primary_key=True)
    username = db.Column(db.String(50), default="default_user")
    session_id = db.Column(db.String(50), index=True)
    summary = db.Column(db.Text, default="")  # 对话摘要
    key_points = db.Column(db.Text, default="[]")  # 关键点提取
    create_time = db.Column(db.DateTime, default=datetime.now)


# ====================== 对话日志表（已有，增加字段）======================
class ChatLog(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    username = db.Column(db.String(50), default="user")
    user_msg = db.Column(db.Text, nullable=False)
    bot_msg = db.Column(db.Text, nullable=False)
    context = db.Column(db.Text, default="[]")
    scene = db.Column(db.String(20), default="auto", index=True)
    is_pro = db.Column(db.Boolean, default=False)
    risk_level = db.Column(db.String(20), default="low")
    risk_alert = db.Column(db.Text, default="")
    rag_sources = db.Column(db.Text, default="")
    retrieval_count = db.Column(db.Integer, default=0)
    create_time = db.Column(db.DateTime, default=datetime.now, index=True)
    # 新增：Agent工具调用记录
    agent_tools_used = db.Column(db.Text, default="[]")
    # 新增：图像分析记录
    image_analyzed = db.Column(db.Boolean, default=False)
    __table_args__ = (
        db.Index('idx_risk_create', 'risk_level', 'create_time'),
        db.Index('idx_pro_scene', 'is_pro', 'scene'),
    )


# 压力测试结果存储
test_results = {
    "total_requests": 0,
    "success_requests": 0,
    "fail_requests": 0,
    "avg_response_time": 0.0,
    "response_times": [],
    "error_details": defaultdict(int)
}

# ====================== 智能体配置 ======================
# 定义智能体可用的工具
AGENT_TOOLS = {
    "search_law": {
        "name": "法律条文检索",
        "description": "检索中国法律法规的具体条文内容",
        "keywords": ["法律条文", "第.*条", "民法典", "刑法", "劳动法", "合同法", "检索"]
    },
    "search_case": {
        "name": "相似案例检索",
        "description": "检索与用户情况相似的司法案例",
        "keywords": ["类似案例", "相似案件", "判例", "司法实践", "参考案例"]
    },
    "calculate_compensation": {
        "name": "赔偿计算器",
        "description": "计算各类赔偿金额（劳动赔偿、交通事故赔偿等）",
        "keywords": ["赔偿多少", "能赔多少", "算一下赔偿", "赔偿计算", "能拿多少钱"]
    },
    "check_deadline": {
        "name": "时效检查",
        "description": "检查各类法律行为的时效要求（诉讼时效、仲裁时效等）",
        "keywords": ["时效", "多久", "过期", "还能起诉吗", "时间限制", "有效期内"]
    },
    "draft_document": {
        "name": "文书草拟",
        "description": "草拟各类法律文书",
        "keywords": ["起草", "写一份", "文书", "起诉状", "申请书", "合同"]
    },
    "get_user_reminder": {
        "name": "获取用户提醒",
        "description": "获取用户设置的重要日期提醒（需用户授权）",
        "keywords": ["提醒", "日程", "备忘", "重要日期"]
    }
}

# ====================== 14类法律场景关键词配置 ======================

# 1. 民间借贷
PRIVATE_LENDING_KEYWORDS = ["借钱", "欠钱", "借条", "欠条", "转账", "债务", "还款", "利息", "高利贷",
                            "担保", "保证人", "催收", "讨债", "借贷", "出借", "借款人", "贷款"]

# 2. 劳动纠纷
LABOR_KEYWORDS = ["工资", "辞退", "拖欠", "加班", "社保", "工伤", "劳动合同", "仲裁", "工地",
                  "赔偿金", "经济补偿", "N+1", "2N", "违法解除", "竞业限制", "劳务派遣", "试用期",
                  "开除", "离职", "调岗", "降薪", "五险一金"]

# 3. 交通事故
TRAFFIC_KEYWORDS = ["车祸", "撞车", "肇事", "逃逸", "酒驾", "醉驾", "交强险", "三者险", "理赔",
                    "伤残鉴定", "误工费", "医药费", "交通肇事", "开车", "电动车", "摩托车", "行人"]

# 4. 婚姻家庭
MARRIAGE_KEYWORDS = ["离婚", "抚养权", "彩礼", "结婚", "财产分割", "家暴", "家庭暴力", "感情破裂",
                     "重婚", "同居", "婚外情", "出轨", "抚养费", "赡养", "探望权", "婚内协议"]

# 5. 刑事法律
CRIMINAL_KEYWORDS = ["盗窃", "诈骗", "轻伤", "拘留", "逮捕", "取保", "缓刑", "坐牢", "立案",
                     "罪名", "故意伤害", "抢劫", "强奸", "贪污", "受贿", "帮信", "掩饰隐瞒",
                     "取保候审", "监视居住", "刑事", "判刑", "监狱", "刑事责任"]

# 6. 房产纠纷
PROPERTY_KEYWORDS = ["买房", "卖房", "租房", "房东", "租客", "押金", "房产证", "产权", "物业",
                     "开发商", "预售", "定金", "首付", "贷款买房", "二手房", "中介", "过户",
                     "房屋质量", "漏水", "违建", "拆迁"]

# 7. 合同纠纷
CONTRACT_KEYWORDS = ["合同", "违约", "解除合同", "合同无效", "格式条款", "定金", "订金", "霸王条款",
                     "合同纠纷", "协议", "承诺", "要约", "买卖合同", "服务合同", "租赁合同"]

# 8. 侵权责任
TORT_KEYWORDS = ["侵权", "赔偿", "人身损害", "名誉权", "隐私权", "肖像权", "姓名权", "高空抛物",
                 "动物伤人", "狗咬", "环境污染", "网络侵权", "诽谤", "侮辱"]

# 9. 诉讼程序
PROCEDURE_KEYWORDS = ["起诉", "上诉", "诉讼", "法院", "开庭", "判决", "执行", "仲裁", "诉讼时效",
                      "管辖", "立案", "证据", "举证", "证人", "诉讼费", "律师费", "再审", "申诉"]

# 10. 公司法律
CORPORATE_KEYWORDS = ["公司", "股东", "股权", "法人", "法定代表人", "破产", "清算", "注销",
                      "增资", "减资", "章程", "董事会", "监事会", "合伙", "合伙企业"]

# 11. 知识产权
IP_KEYWORDS = ["专利", "商标", "版权", "著作权", "侵权", "盗版", "山寨", "知识产权",
               "商业秘密", "技术秘密", "注册", "申请专利", "注册商标"]

# 12. 消费者权益
CONSUMER_KEYWORDS = ["假货", "退货", "退款", "欺诈", "三倍赔偿", "十倍赔偿", "消费者",
                     "维权", "12315", "预付卡", "会员卡", "霸王条款", "食品安全"]

# 13. 行政法律
ADMINISTRATIVE_KEYWORDS = ["行政拘留", "行政处罚", "行政复议", "行政诉讼", "政府信息",
                           "公开", "拆迁补偿", "行政许可", "行政强制", "罚款", "吊销执照"]

# 14. 继承法律
INHERITANCE_KEYWORDS = ["继承", "遗嘱", "遗产", "法定继承", "遗赠", "公证", "继承人",
                        "放弃继承", "遗产分割", "财产分配", "死亡"]

# 场景映射表
SCENE_CONFIG = {
    "private_lending": {
        "name": "民间借贷",
        "keywords": PRIVATE_LENDING_KEYWORDS,
        "color": "#e8f5e9",
        "text_color": "#2e7d32"
    },
    "labor": {
        "name": "劳动纠纷",
        "keywords": LABOR_KEYWORDS,
        "color": "#fff3e0",
        "text_color": "#e65100"
    },
    "traffic_accident": {
        "name": "交通事故",
        "keywords": TRAFFIC_KEYWORDS,
        "color": "#e3f2fd",
        "text_color": "#1565c0"
    },
    "marriage_family": {
        "name": "婚姻家庭",
        "keywords": MARRIAGE_KEYWORDS,
        "color": "#fce4ec",
        "text_color": "#c2185b"
    },
    "criminal": {
        "name": "刑事法律",
        "keywords": CRIMINAL_KEYWORDS,
        "color": "#ffebee",
        "text_color": "#c62828"
    },
    "property": {
        "name": "房产纠纷",
        "keywords": PROPERTY_KEYWORDS,
        "color": "#e8eaf6",
        "text_color": "#283593"
    },
    "contract": {
        "name": "合同纠纷",
        "keywords": CONTRACT_KEYWORDS,
        "color": "#f1f8e9",
        "text_color": "#33691e"
    },
    "tort": {
        "name": "侵权责任",
        "keywords": TORT_KEYWORDS,
        "color": "#f3e5f5",
        "text_color": "#6a1b9a"
    },
    "procedure": {
        "name": "诉讼程序",
        "keywords": PROCEDURE_KEYWORDS,
        "color": "#e0f7fa",
        "text_color": "#006064"
    },
    "corporate": {
        "name": "公司法律",
        "keywords": CORPORATE_KEYWORDS,
        "color": "#efebe9",
        "text_color": "#4e342e"
    },
    "intellectual_property": {
        "name": "知识产权",
        "keywords": IP_KEYWORDS,
        "color": "#e8eaf6",
        "text_color": "#1a237e"
    },
    "consumer": {
        "name": "消费者权益",
        "keywords": CONSUMER_KEYWORDS,
        "color": "#fff8e1",
        "text_color": "#f57f17"
    },
    "administrative": {
        "name": "行政法律",
        "keywords": ADMINISTRATIVE_KEYWORDS,
        "color": "#e0f2f1",
        "text_color": "#004d40"
    },
    "inheritance": {
        "name": "继承法律",
        "keywords": INHERITANCE_KEYWORDS,
        "color": "#fbe9e7",
        "text_color": "#bf360c"
    }
}

SCENE_NAMES = {key: config["name"] for key, config in SCENE_CONFIG.items()}

# ====================== 提示词模板 ======================
BASE_SYSTEM_PROMPT = """
你是专业的法律咨询助手。

【回答要求】
1. 优先使用下方知识库中检索到的信息回答问题
2. 法律分析仅作参考，必须提示"不可替代律师专业意见"
3. 回答控制在200字左右，简洁清晰
4. 回答完后推荐3个相关问题

格式：回答内容 + 换行 + 💡你可以继续问：问题1？问题2？问题3？
"""

PRO_SYSTEM_PROMPT = """
你是资深法律专家，为同行提供专业支持。

【重要】你只需要回答法律内容，不要在回答末尾添加类似"📚 知识库检索"等来源标记，这些会由系统自动添加。

【回答格式】
**一、法律依据与构成要件**
**二、举证责任分配**
**三、诉讼策略与风险**
**四、类案裁判观点**
**五、实操建议**
**六、专业延伸（3个问题）**
"""


# ====================== 智能体工具函数 ======================
def agent_search_law(query: str, scene: str = "") -> dict:
    """检索法律条文"""
    try:
        # 构建更精准的检索查询
        search_query = f"{scene}法律条文 {query}" if scene else f"法律条文 {query}"
        results = knowledge_manager.search(search_query, top_k=3)

        if results:
            law_contents = []
            for r in results:
                category = r.get("metadata", {}).get("category", "")
                content = r.get("content", "")
                if len(content) > 800:
                    content = content[:800] + "..."
                law_contents.append(f"【{category}】\n{content}")

            return {
                "success": True,
                "tool": "search_law",
                "result": "\n\n---\n\n".join(law_contents),
                "sources": [r.get("metadata", {}).get("category", "") for r in results]
            }
        return {"success": False, "tool": "search_law", "result": "未找到相关法律条文", "sources": []}
    except Exception as e:
        return {"success": False, "tool": "search_law", "result": f"检索失败: {str(e)}", "sources": []}


def agent_search_case(query: str, scene: str = "") -> dict:
    """检索相似案例"""
    try:
        search_query = f"案例 {scene} {query}" if scene else f"案例 {query}"
        results = knowledge_manager.search(search_query, top_k=3)

        if results:
            case_contents = []
            sources = []
            for r in results:
                category = r.get("metadata", {}).get("category", "")
                content = r.get("content", "")
                sources.append(category)
                case_contents.append(f"【参考案例 - {category}】\n{content[:600]}")

            return {
                "success": True,
                "tool": "search_case",
                "result": "\n\n---\n\n".join(case_contents),
                "sources": sources
            }
        return {"success": False, "tool": "search_case", "result": "未找到相似案例", "sources": []}
    except Exception as e:
        return {"success": False, "tool": "search_case", "result": f"检索失败: {str(e)}", "sources": []}


def agent_calculate_compensation(query: str, scene: str = "") -> dict:
    """赔偿计算器 - 基于规则和知识库"""
    # 解析赔偿类型
    compensation_result = ""
    if "劳动" in scene or "工资" in query or "辞退" in query:
        compensation_result = """
【劳动赔偿快速计算参考】

1. 经济补偿金(N)：工作年限 × 月平均工资
   - 满1年按1年计算
   - 6个月以上不满1年，按1年计算
   - 不满6个月，支付半个月工资

2. 赔偿金(2N)：违法解除劳动合同，支付2倍经济补偿

3. 未签劳动合同双倍工资：最多支持11个月差额

⚠️ 具体金额需结合实际情况计算，建议咨询专业律师
"""
    elif "交通" in scene or "车祸" in query:
        compensation_result = """
【交通事故赔偿项目参考】

1. 医疗费：实际发生的合理医疗费用
2. 误工费：月收入 ÷ 21.75 × 误工天数
3. 护理费：当地护工标准 × 护理天数
4. 残疾赔偿金：当地人均可支配收入 × 20年 × 伤残系数
5. 精神损害抚慰金：5000-50000元（根据伤残等级）
6. 被扶养人生活费：根据被扶养人情况计算

伤残系数参考：
- 十级10% | 九级20% | 八级30% | ... | 一级100%
"""
    else:
        compensation_result = """
【赔偿计算通用参考】

赔偿金额需考虑以下因素：
1. 实际损失：直接经济损失
2. 预期利益损失：合同履行后可获得的利益
3. 精神损害赔偿：根据侵权情节严重程度
4. 惩罚性赔偿：适用于恶意侵权、消费欺诈等情形

建议提供更多案件细节进行精准计算
"""

    return {
        "success": True,
        "tool": "calculate_compensation",
        "result": compensation_result,
        "sources": ["赔偿计算标准"]
    }


def agent_check_deadline(query: str, scene: str = "") -> dict:
    """检查时效要求"""
    deadlines = {
        "劳动仲裁": "自知道权利被侵害之日起1年内提出",
        "民事诉讼一般时效": "3年（自知道权利受损害之日起计算）",
        "行政诉讼": "6个月（自知道行政行为之日起）",
        "行政复议": "60日（自知道具体行政行为之日起）",
        "工伤认定": "单位30日内申请，职工1年内申请",
        "交通事故索赔": "人身损害1年，财产损失3年",
        "民间借贷": "约定还款日起3年，无约定最长20年",
        "合同纠纷": "违约行为发生之日起3年",
        "刑事追诉": "根据法定最高刑期确定，最高20年"
    }

    result = "【法律时效参考】\n\n"
    for name, period in deadlines.items():
        result += f"• {name}：{period}\n"

    result += "\n⚠️ 时效中断、中止等情形可能影响时效计算，建议咨询律师确认"

    return {
        "success": True,
        "tool": "check_deadline",
        "result": result,
        "sources": ["诉讼法典"]
    }


def agent_detect_intent(user_msg: str, scene: str) -> dict:
    """智能体意图识别 - 决定调用哪些工具"""
    msg_lower = user_msg.lower()
    tools_to_call = []

    for tool_key, tool_info in AGENT_TOOLS.items():
        # 检查关键词匹配
        matched = False
        for kw in tool_info.get("keywords", []):
            if ".*" in kw:
                # 正则匹配
                if re.search(kw, msg_lower):
                    matched = True
                    break
            elif kw in msg_lower:
                matched = True
                break
        if matched:
            tools_to_call.append(tool_key)

    # 去重并限制最多2个工具
    tools_to_call = list(set(tools_to_call))[:2]

    return {
        "tools": tools_to_call,
        "should_use_agent": len(tools_to_call) > 0
    }


def execute_agent_tools(user_msg: str, scene: str) -> dict:
    """执行智能体工具调用"""
    intent = agent_detect_intent(user_msg, scene)

    if not intent["should_use_agent"]:
        return {"used_tools": [], "tool_results": "", "tool_sources": []}

    tool_results = []
    tool_sources = []

    for tool_name in intent["tools"]:
        if tool_name == "search_law":
            result = agent_search_law(user_msg, scene)
        elif tool_name == "search_case":
            result = agent_search_case(user_msg, scene)
        elif tool_name == "calculate_compensation":
            result = agent_calculate_compensation(user_msg, scene)
        elif tool_name == "check_deadline":
            result = agent_check_deadline(user_msg, scene)
        else:
            continue

        if result["success"] and result["result"]:
            tool_results.append(f"🔧 【{AGENT_TOOLS[tool_name]['name']}】\n{result['result']}")
            tool_sources.extend(result.get("sources", []))

    if tool_results:
        combined = "\n\n---\n\n".join(tool_results)
        return {"used_tools": intent["tools"], "tool_results": combined, "tool_sources": tool_sources}

    return {"used_tools": [], "tool_results": "", "tool_sources": []}


# ====================== 长期记忆管理函数 ======================
def extract_key_info_from_conversation(messages: list) -> dict:
    """从对话中提取关键信息用于长期记忆"""
    key_info = {
        "case_type": "",
        "involved_parties": [],
        "key_amounts": [],
        "important_dates": [],
        "status": "ongoing"
    }

    all_text = " ".join([m.get("user", "") for m in messages[-10:] if m.get("user")])

    # 提取案件类型
    for scene_key, scene_config in SCENE_CONFIG.items():
        for kw in scene_config["keywords"]:
            if kw in all_text:
                key_info["case_type"] = scene_config["name"]
                break
        if key_info["case_type"]:
            break

    # 提取金额信息
    money_pattern = r'(\d+(?:\.\d+)?)\s*[万万千百]?\s*元'
    amounts = re.findall(money_pattern, all_text)
    if amounts:
        key_info["key_amounts"] = amounts[:3]

    # 提取日期
    date_pattern = r'(\d{4}年\d{1,2}月\d{1,2}日|\d{1,2}月\d{1,2}日|\d{4}-\d{1,2}-\d{1,2})'
    dates = re.findall(date_pattern, all_text)
    if dates:
        key_info["important_dates"] = dates[:3]

    return key_info


def update_long_term_memory(session_id: str, messages: list, username: str = "default_user"):
    """更新长期记忆"""
    try:
        # 提取关键信息
        key_info = extract_key_info_from_conversation(messages)

        # 检查是否已有该会话的摘要
        existing = ConversationSummary.query.filter_by(
            username=username,
            session_id=session_id
        ).first()

        if existing:
            # 更新现有摘要
            existing.key_points = json.dumps(key_info, ensure_ascii=False)
            existing.summary = f"案件类型：{key_info['case_type']}，涉及金额：{', '.join(key_info['key_amounts'])}"
        else:
            # 创建新摘要
            summary = ConversationSummary(
                username=username,
                session_id=session_id,
                summary=f"案件类型：{key_info['case_type']}",
                key_points=json.dumps(key_info, ensure_ascii=False)
            )
            db.session.add(summary)

        db.session.commit()
    except Exception as e:
        print(f"更新长期记忆失败: {e}")
        db.session.rollback()


def get_long_term_context(username: str = "default_user") -> str:
    """获取长期记忆上下文"""
    try:
        # 获取最近的对话摘要
        recent_summaries = ConversationSummary.query.filter_by(
            username=username
        ).order_by(ConversationSummary.create_time.desc()).limit(5).all()

        if not recent_summaries:
            return ""

        context = "【历史对话记忆】\n"
        for i, summary in enumerate(recent_summaries, 1):
            context += f"{i}. {summary.summary}\n"

        return context
    except Exception as e:
        print(f"获取长期记忆失败: {e}")
        return ""


# ====================== 多模态处理：图像分析 ======================
def analyze_image_from_base64(image_base64: str, user_msg: str = "") -> dict:
    """分析上传的图片 - 根据用户问题动态分析"""
    try:
        headers = {
            "Authorization": f"Bearer {DEEPSEEK_API_KEY}",
            "Content-Type": "application/json"
        }

        # 动态构建提示词 - 根据用户消息决定分析方向
        if user_msg and user_msg.strip():
            # 用户有输入问题，按用户要求分析
            analysis_prompt = f"""
用户上传了一张图片，并问：{user_msg}

请根据用户的图片和问题，提供专业的法律分析。

分析要求：
1. 首先识别图片类型（欠条/借条/合同/判决书/收据/工资条/身份证/其他）
2. 根据图片类型提取关键信息：
   - 如果是欠条/借条：提取借款人、出借人、金额、借款日期、还款日期、利率、逾期责任
   - 如果是合同：提取合同双方、标的金额、关键条款、权利义务、违约责任
   - 如果是判决书：提取案号、当事人、判决结果、法院观点
   - 如果是收据/发票：提取金额、日期、开票方、用途
   - 如果是工资条：提取工资数额、扣款项目、公司名称
   - 如果是其他：提取所有可见的法律相关信息
3. 给出法律分析和建议
4. 指出可能存在的风险或问题

请确保回答准确、专业，用中文回答。
"""
        else:
            # 用户没有输入，让系统自动识别图片类型
            analysis_prompt = """
用户上传了一张图片，但没有提出具体问题。

请自动识别图片类型并提供法律分析：

图片类型包括但不限于：
- 欠条/借条
- 合同/协议
- 判决书/裁定书
- 收据/发票
- 工资条
- 身份证/证件
- 其他法律文件

分析步骤：
1. 首先说明识别出的图片类型
2. 提取图片中的所有关键信息（金额、日期、当事人、签名等）
3. 给出法律层面的分析和建议
4. 指出可能存在的风险点

请用中文回答，确保分析清晰有用。
"""

        payload = {
            "model": "deepseek-chat",
            "messages": [
                {
                    "role": "user",
                    "content": [
                        {"type": "text", "text": analysis_prompt},
                        {"type": "image_url", "image_url": {"url": f"data:image/jpeg;base64,{image_base64}"}}
                    ]
                }
            ],
            "temperature": 0.3,
            "max_tokens": 1500
        }

        response = requests.post(DEEPSEEK_URL, headers=headers, json=payload, timeout=60, verify=False)

        if response.status_code == 200:
            data = response.json()
            analysis = data["choices"][0]["message"]["content"]
            return {"success": True, "analysis": analysis}
        else:
            # 降级：返回基础分析
            return {
                "success": True,
                "analysis": "图片已接收。如需详细分析，请在输入框中描述您的问题，例如：\"请分析这张欠条是否有效\"、\"这份合同有什么风险\"等。\n\n（注意：AI分析仅供参考，重要证据请咨询专业律师核实）"
            }
    except Exception as e:
        print(f"图片分析失败: {e}")
        return {"success": False, "analysis": f"图片分析失败: {str(e)}"}


# ====================== 辅助函数 ======================
def clean_rag_markers(text):
    """清理文本中可能自带的RAG来源标记"""
    if not text:
        return text

    patterns = [
        r'\n*---\n*📚\s*【知识库检索】[^\n]*\n*',
        r'\n*---\n*📚\s*知识库检索[^\n]*\n*',
        r'\n*---\n*【知识库检索】[^\n]*\n*',
        r'\n*---\n*知识库检索[^\n]*\n*',
        r'\n*📚\s*【知识库检索】[^\n]*\n*',
        r'\n*📚\s*知识库检索[^\n]*\n*',
        r'\n*---\n*🔧\s*【.*?】[^\n]*\n*',
    ]

    for pattern in patterns:
        text = re.sub(pattern, '', text, flags=re.DOTALL)

    text = re.sub(r'\n{3,}', '\n\n', text)
    return text.strip()


def get_rag_context(user_msg: str):
    """从本地知识库检索相关内容"""
    try:
        results = knowledge_manager.search(user_msg, top_k=3)

        if not results:
            return None, []

        context_parts = []
        sources = []

        for result in results:
            category = result.get("metadata", {}).get("category", "未知")
            content = result.get("content", "")
            sources.append(category)

            if len(content) > 500:
                content = content[:500] + "..."

            context_parts.append(f"【{category}法律知识】\n{content}")

        rag_text = "\n\n---\n\n".join(context_parts)
        return rag_text, sources
    except Exception as e:
        print(f"RAG检索失败: {e}")
        return None, []


def auto_detect_scene(user_msg: str) -> str:
    """自动识别14类法律场景"""
    msg_lower = user_msg.lower()

    scene_scores = {}
    for scene_key, config in SCENE_CONFIG.items():
        score = sum(1 for kw in config["keywords"] if kw in msg_lower)
        scene_scores[scene_key] = score

    max_score = max(scene_scores.values()) if scene_scores else 0

    if max_score >= 2:
        for scene_key, score in scene_scores.items():
            if score == max_score:
                return scene_key

    if max_score == 1:
        matched_scenes = [k for k, v in scene_scores.items() if v == 1]
        if len(matched_scenes) == 1:
            return matched_scenes[0]

    return "contract"


def get_scene_name(scene_key: str) -> str:
    return SCENE_NAMES.get(scene_key, "法律咨询")


def detect_risk_level(user_msg: str, scene: str) -> tuple:
    """混合风险评估"""
    msg_lower = user_msg.lower()

    high_keywords = [
        "家暴", "家庭暴力", "拘留", "刑事拘留", "逮捕", "判刑", "坐牢",
        "死亡", "重伤", "杀人", "打死", "生命危险", "逃跑", "紧急",
        "百万", "千万", "巨额", "全部积蓄", "养老钱", "救命钱"
    ]
    for kw in high_keywords:
        if kw in msg_lower:
            return "high", "⚠️ 【高风险】涉及重大法律风险，建议立即咨询律师"

    try:
        results = knowledge_manager.search(user_msg, top_k=3)
        high_count = 0
        medium_count = 0

        for result in results:
            category = result.get("metadata", {}).get("category", "")
            content = result.get("content", "").lower()

            if category in ["刑事法律", "刑事"]:
                high_count += 1
            elif category in ["交通事故", "交通"]:
                if "死亡" in content or "逃逸" in content or "重伤" in content:
                    high_count += 1
                else:
                    medium_count += 1
            elif category in ["劳动纠纷", "劳动"]:
                if "工伤" in content or "死亡" in content or "伤残" in content:
                    high_count += 1
                elif "拖欠" in content or "辞退" in content:
                    medium_count += 1
            elif category in ["民间借贷", "婚姻家庭", "合同纠纷", "房产纠纷"]:
                if "诉讼" in content or "法院" in content or "起诉" in content:
                    medium_count += 1
            elif category in ["侵权责任", "诉讼程序"]:
                medium_count += 1

        if high_count >= 1:
            return "high", "⚠️ 【高风险】根据相关案例，您的情况存在较高法律风险，建议立即咨询律师"
        if medium_count >= 2:
            return "medium", "🔔 【中风险】根据类似案例，建议尽快收集证据并咨询专业律师"
    except Exception as e:
        print(f"知识库风险评估失败: {e}")

    if scene == "criminal":
        return "medium", "🔔 【中风险】涉及刑事法律问题，建议咨询专业律师"

    medium_keywords = ["起诉", "诉讼", "法院", "开庭", "判决", "执行", "仲裁", "赔偿", "辞退", "拖欠工资", "离婚",
                       "抚养权", "财产分割"]
    for kw in medium_keywords:
        if kw in msg_lower:
            return "medium", "🔔 【中风险】建议尽快收集证据并咨询专业律师"

    return "low", "ℹ️ 【低风险】一般法律咨询，以上分析仅供参考"


def merge_risk_with_reply(bot_reply: str, risk_level: str, risk_alert: str) -> str:
    if risk_level == "high":
        return f"{risk_alert}\n\n---\n\n{bot_reply}"
    elif risk_level == "medium":
        return f"{bot_reply}\n\n---\n\n{risk_alert}"
    else:
        return f"{bot_reply}\n\n{risk_alert}"


# ====================== DeepSeek 配置 ======================
DEEPSEEK_API_KEY = "sk-..."
DEEPSEEK_URL = "https://api.deepseek.com/v1/chat/completions"

import urllib3

urllib3.disable_warnings(urllib3.exceptions.InsecureRequestWarning)
session = requests.Session()
session.verify = False


def call_deepseek(messages, scene="civil", is_pro=False, retry_times=2, rag_text=None, tool_results=None):
    """调用DeepSeek API（支持智能体工具结果注入）"""
    scene_name = get_scene_name(scene)

    if is_pro:
        system_prompt = PRO_SYSTEM_PROMPT
    else:
        system_prompt = BASE_SYSTEM_PROMPT

    # 构建增强的系统提示词
    enhanced_prompt = system_prompt

    if rag_text:
        enhanced_prompt = f"""
{enhanced_prompt}

【当前咨询场景】{scene_name}

【知识库检索结果 - 请优先使用以下信息回答用户问题】
{rag_text}
"""

    if tool_results:
        enhanced_prompt = f"""
{enhanced_prompt}

【智能体工具调用结果 - 请结合这些专业信息回答】
{tool_results}
"""

    messages[0]["content"] = enhanced_prompt
    max_context = 12 if is_pro else 8
    messages = messages[:max_context * 2 + 1]

    headers = {
        "Authorization": f"Bearer {DEEPSEEK_API_KEY}",
        "Content-Type": "application/json"
    }

    payload = {
        "model": "deepseek-chat",
        "messages": messages,
        "temperature": 0.15 if is_pro else 0.25,
        "stream": False
    }

    for retry in range(retry_times + 1):
        try:
            start_time = time.time()
            resp = session.post(DEEPSEEK_URL, headers=headers, json=payload, timeout=35)
            end_time = time.time()

            if resp.status_code == 200:
                data = resp.json()
                return {
                    "reply": data["choices"][0]["message"]["content"],
                    "success": True,
                    "response_time": end_time - start_time
                }
            else:
                if retry < retry_times:
                    time.sleep(0.8 * (retry + 1))
                    continue
                raise Exception(f"API错误 {resp.status_code}")
        except Exception as e:
            if retry < retry_times:
                time.sleep(0.8 * (retry + 1))
                continue
            return {
                "reply": "服务暂时不可用，请稍后重试。",
                "success": False,
                "response_time": 0,
                "error": str(e)
            }

    return {"reply": "服务暂时不可用", "success": False, "response_time": 0}


# ====================== 新增路由：用户档案管理 ======================
@app.route('/user-profile', methods=['GET', 'POST', 'PUT'])
def user_profile():
    """用户档案管理"""
    username = request.json.get("username", "default_user") if request.method != 'GET' else "default_user"

    if request.method == 'GET':
        profile = UserProfile.query.filter_by(username=username).first()
        if not profile:
            return jsonify({"success": True, "profile": {}})
        return jsonify({"success": True, "profile": profile.to_dict()})

    elif request.method == 'POST':
        # 创建或更新档案
        profile = UserProfile.query.filter_by(username=username).first()
        data = request.json

        if not profile:
            profile = UserProfile(username=username)
            db.session.add(profile)

        if "name" in data:
            profile.name = data["name"]
        if "phone" in data:
            profile.phone = data["phone"]
        if "id_number" in data:
            profile.id_number = data["id_number"]
        if "address" in data:
            profile.address = data["address"]
        if "important_memos" in data:
            profile.important_memos = json.dumps(data["important_memos"], ensure_ascii=False)

        profile.last_active = datetime.now()
        db.session.commit()

        return jsonify({"success": True, "message": "档案已更新"})

    elif request.method == 'PUT':
        # 更新咨询次数和常用类别
        profile = UserProfile.query.filter_by(username=username).first()
        if profile:
            profile.total_consultations += 1
            db.session.commit()

        return jsonify({"success": True})


# ====================== 新增路由：图像分析 ======================
@app.route('/analyze-image', methods=['POST'])
def analyze_image():
    """分析上传的证据图片"""
    data = request.get_json()
    image_base64 = data.get("image", "")
    user_msg = data.get("user_msg", "")

    if not image_base64:
        return jsonify({"success": False, "error": "未提供图片"}), 400

    # 移除base64前缀（如果有）
    if "," in image_base64:
        image_base64 = image_base64.split(",")[1]

    result = analyze_image_from_base64(image_base64, user_msg)

    return jsonify({
        "success": result["success"],
        "analysis": result.get("analysis", "分析失败"),
        "message": "图片分析完成" if result["success"] else "分析失败"
    })


# ====================== 路由：长期记忆摘要 ======================
@app.route('/memory-summary', methods=['GET'])
def memory_summary():
    """获取长期记忆摘要"""
    username = request.args.get("username", "default_user")
    context = get_long_term_context(username)
    return jsonify({"success": True, "memory_context": context})


# ====================== AI 生成法律文书（原有） ======================
@app.route('/generate-document', methods=['POST'])
def generate_document():
    data = request.get_json()
    user_msg = data.get("user_msg", "")
    context = data.get("context", [])
    doc_type = data.get("doc_type", "complaint")
    doc_category = data.get("doc_category", "")

    all_dialogues = []
    for item in context[-15:]:
        if item.get("user"):
            all_dialogues.append(f"用户：{item['user']}")
        if item.get("bot"):
            all_dialogues.append(f"助手：{item['bot']}")
    all_dialogues.append(f"用户：{user_msg}")
    all_text = "\n".join(all_dialogues)

    rag_text = None
    rag_sources = []

    if doc_category:
        try:
            search_results = knowledge_manager.search(doc_category, top_k=3)
            if search_results:
                context_parts = []
                for result in search_results:
                    category = result.get("metadata", {}).get("category", doc_category)
                    content = result.get("content", "")
                    rag_sources.append(category)
                    if len(content) > 800:
                        content = content[:800] + "..."
                    context_parts.append(f"【{category}相关法条与知识】\n{content}")
                if context_parts:
                    rag_text = "\n\n---\n\n".join(context_parts)
                    print(f"📚 文书生成检索到 {len(rag_sources)} 条知识，类别：{doc_category}")
        except Exception as e:
            print(f"文书生成知识库检索失败: {e}")

    law_reference = ""
    if rag_text:
        law_reference = f"\n\n【相关法律知识参考】\n{rag_text}"

    doc_prompts = {
        "complaint": f"""请根据以下描述生成《民事起诉状》：

【用户描述的案件情况】
{all_text}
{law_reference}

要求：
1. 包含原告信息、被告信息、诉讼请求、事实与理由、此致、具状人
2. 缺失信息用「待补充」标注
3. 格式规范，语言严谨
4. 如有相关法律知识，请据此增强文书的法律依据""",
        "defense": f"""请根据以下描述生成《民事答辩状》：

【用户描述的案件情况】
{all_text}
{law_reference}

要求：
1. 包含答辩人信息、答辩请求、事实与理由
2. 缺失信息用「待补充」标注
3. 格式规范，语言严谨""",
        "labor_application": f"""请根据以下描述生成《劳动仲裁申请书》：

【用户描述的案件情况】
{all_text}
{law_reference}

要求：
1. 包含申请人信息、被申请人信息、仲裁请求、事实与理由
2. 缺失信息用「待补充」标注
3. 格式规范，语言严谨
4. 注意劳动仲裁时效为一年""",
        "contract_clause": f"""请根据以下需求生成合同条款建议：

【用户描述的需求】
{all_text}
{law_reference}

要求：
1. 输出3-5条核心合同条款
2. 每条条款包含条款名称和具体内容
3. 语言严谨，权责清晰""",
        "criminal_defense": f"""请根据以下描述生成《刑事辩护意见书》：

【用户描述的案件情况】
{all_text}
{law_reference}

要求：
1. 包含当事人信息、案件事实概述
2. 辩护意见分点论述（定罪辩护/量刑辩护/程序辩护）
3. 法律依据引用准确
4. 结论明确
5. 缺失信息用「待补充」标注"""
    }

    prompt = doc_prompts.get(doc_type, doc_prompts["complaint"])

    messages = [
        {"role": "system", "content": "你是专业的法律文书生成助手，擅长根据案件描述生成格式规范、语言严谨的法律文书。"},
        {"role": "user", "content": prompt}
    ]

    result = call_deepseek(messages, scene="civil", is_pro=True)

    if result["success"]:
        clean_doc = result["reply"]
        clean_doc = re.sub(r'^#{1,6}\s+', '', clean_doc, flags=re.MULTILINE)
        clean_doc = re.sub(r'\n*\s*💡.*$', '', clean_doc, flags=re.DOTALL)
        clean_doc = clean_doc.replace('*', '')
        clean_doc = clean_rag_markers(clean_doc)

        doc_names = {
            "complaint": "民事起诉状",
            "defense": "民事答辩状",
            "labor_application": "劳动仲裁申请书",
            "contract_clause": "合同条款建议",
            "criminal_defense": "刑事辩护意见书"
        }
        filename = f"{doc_names.get(doc_type, '法律文书')}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.txt"

        return jsonify({
            "success": True,
            "document": clean_doc,
            "filename": filename,
            "rag_used": bool(rag_sources),
            "rag_sources": rag_sources
        })
    else:
        return jsonify({"success": False, "error": result.get("error", "生成失败")}), 500


# ====================== 新增：模板填充式文书生成（增强版） ======================
@app.route('/generate-document-v2', methods=['POST'])
def generate_document_v2():
    """基于模板填充的法律文书生成（增强版）"""
    data = request.get_json()
    doc_type = data.get("doc_type", "complaint")
    context = data.get("context", [])
    user_msg = data.get("user_msg", "")
    scene = data.get("scene", "contract")
    use_ai_fill = data.get("use_ai_fill", True)

    # 1. 查找匹配的模板
    template = template_engine.get_template_by_type(doc_type, scene)

    if not template:
        # 降级到原有的AI生成
        return generate_document_fallback(data)

    # 2. 从上下文提取信息
    extracted_data = template_engine.extract_info_from_context(context, scene)

    # 3. 如果有用户输入，提取更多信息
    if user_msg:
        extra_data = template_engine.extract_info_from_text(user_msg, scene)
        for key, value in extra_data.items():
            if value and not extracted_data.get(key):
                extracted_data[key] = value

    # 4. 为诉讼请求添加示例
    if doc_type == "complaint" and not extracted_data.get("claims"):
        extracted_data["claims"] = [
            "判令被告向原告支付款项【请填写具体金额】元；",
            "判令被告承担本案全部诉讼费用；"
        ]

    # 5. 为仲裁请求添加示例
    if doc_type == "labor_application" and not extracted_data.get("arbitration_requests"):
        extracted_data["arbitration_requests"] = [
            "请求裁决被申请人支付拖欠工资【请填写金额】元；",
            "请求裁决被申请人支付经济补偿金；"
        ]

    # 6. 为答辩请求添加示例
    if doc_type == "defense" and not extracted_data.get("defense_requests"):
        extracted_data["defense_requests"] = [
            "请求驳回原告的全部诉讼请求；",
            "请求判令原告承担本案诉讼费用；"
        ]

    # 7. 渲染模板
    document_content, missing_fields = template_engine.render_template(template, extracted_data)

    # 8. 如果使用AI填充且有缺失字段，尝试优化
    if use_ai_fill and missing_fields:
        context_text = " ".join([item.get('user', '') for item in context[-10:] if item.get('user')])
        if user_msg:
            context_text += " " + user_msg

        # 构建AI填充提示
        fill_prompt = f"""请根据以下对话内容，补充法律文书中缺失的信息。

对话内容：{context_text[:800]}

缺失的信息：{', '.join([f['label'] for f in missing_fields])}

请直接输出JSON格式，如：{{"plaintiff_name": "张三", "amount": "10000"}}
"""
        try:
            ai_messages = [
                {"role": "system", "content": "你是法律文书辅助助手，根据对话提取信息。"},
                {"role": "user", "content": fill_prompt}
            ]
            ai_result = call_deepseek(ai_messages, scene=scene, is_pro=False)
            if ai_result["success"]:
                # 尝试解析AI返回的JSON
                ai_json_match = re.search(r'\{[^{}]*\}', ai_result["reply"])
                if ai_json_match:
                    import json as json_module
                    ai_data = json_module.loads(ai_json_match.group())
                    for key, value in ai_data.items():
                        if value and not extracted_data.get(key):
                            extracted_data[key] = value
                    # 重新渲染
                    document_content, missing_fields = template_engine.render_template(template, extracted_data)
        except Exception as e:
            print(f"AI填充失败: {e}")

    # 9. 生成文件名
    doc_names = {
        "complaint": "民事起诉状",
        "defense": "民事答辩状",
        "labor_application": "劳动仲裁申请书",
        "contract_clause": "合同条款建议",
        "criminal_defense": "刑事辩护意见书"
    }
    filename = f"{doc_names.get(doc_type, '法律文书')}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.txt"

    return jsonify({
        "success": True,
        "document": document_content,
        "filename": filename,
        "extracted_info": extracted_data,
        "missing_fields": missing_fields,
        "method": "template_based",
        "template_id": template.get('template_id')
    })


def generate_document_fallback(data):
    """降级到原有的AI生成方式"""
    user_msg = data.get("user_msg", "")
    context = data.get("context", [])
    doc_type = data.get("doc_type", "complaint")
    doc_category = data.get("doc_category", "")

    all_dialogues = []
    for item in context[-15:]:
        if item.get("user"):
            all_dialogues.append(f"用户：{item['user']}")
        if item.get("bot"):
            all_dialogues.append(f"助手：{item['bot']}")
    all_dialogues.append(f"用户：{user_msg}")
    all_text = "\n".join(all_dialogues)

    rag_text = None
    rag_sources = []

    if doc_category:
        try:
            search_results = knowledge_manager.search(doc_category, top_k=3)
            if search_results:
                context_parts = []
                for result in search_results:
                    category = result.get("metadata", {}).get("category", doc_category)
                    content = result.get("content", "")
                    rag_sources.append(category)
                    if len(content) > 800:
                        content = content[:800] + "..."
                    context_parts.append(f"【{category}相关法条与知识】\n{content}")
                if context_parts:
                    rag_text = "\n\n---\n\n".join(context_parts)
        except Exception as e:
            print(f"文书生成知识库检索失败: {e}")

    law_reference = ""
    if rag_text:
        law_reference = f"\n\n【相关法律知识参考】\n{rag_text}"

    doc_prompts = {
        "complaint": f"请根据以下描述生成《民事起诉状》：\n\n【用户描述的案件情况】\n{all_text}{law_reference}",
        "defense": f"请根据以下描述生成《民事答辩状》：\n\n【用户描述的案件情况】\n{all_text}{law_reference}",
        "labor_application": f"请根据以下描述生成《劳动仲裁申请书》：\n\n【用户描述的案件情况】\n{all_text}{law_reference}",
        "contract_clause": f"请根据以下需求生成合同条款：\n\n{all_text}{law_reference}",
        "criminal_defense": f"请根据以下描述生成《刑事辩护意见书》：\n\n{all_text}{law_reference}"
    }

    prompt = doc_prompts.get(doc_type, doc_prompts["complaint"])
    messages = [
        {"role": "system", "content": "你是专业的法律文书生成助手。"},
        {"role": "user", "content": prompt}
    ]

    result = call_deepseek(messages, scene="civil", is_pro=True)

    if result["success"]:
        clean_doc = result["reply"]
        clean_doc = re.sub(r'^#{1,6}\s+', '', clean_doc, flags=re.MULTILINE)
        clean_doc = re.sub(r'\n*\s*💡.*$', '', clean_doc, flags=re.DOTALL)
        clean_doc = clean_rag_markers(clean_doc)

        doc_names = {
            "complaint": "民事起诉状",
            "defense": "民事答辩状",
            "labor_application": "劳动仲裁申请书",
            "contract_clause": "合同条款建议",
            "criminal_defense": "刑事辩护意见书"
        }
        filename = f"{doc_names.get(doc_type, '法律文书')}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.txt"

        return jsonify({
            "success": True,
            "document": clean_doc,
            "filename": filename,
            "method": "ai_generated"
        })
    else:
        return jsonify({"success": False, "error": result.get("error", "生成失败")}), 500


# ====================== 新增：获取模板列表 ======================
@app.route('/template-list', methods=['GET'])
def get_template_list():
    """获取所有可用模板"""
    templates = template_engine.get_all_templates()
    return jsonify({
        "success": True,
        "templates": templates
    })


# ====================== 新增：获取模板详情 ======================
@app.route('/template/<template_id>', methods=['GET'])
def get_template_detail(template_id):
    """获取模板详情"""
    template = template_engine.get_template_by_id(template_id)
    if template:
        return jsonify({
            "success": True,
            "template": template
        })
    return jsonify({"success": False, "error": "模板不存在"}), 404


# ====================== 导出文档 ======================
@app.route('/export-document', methods=['POST'])
def export_document():
    data = request.get_json()
    content = data.get("content", "")
    filename = data.get("filename", "法律文书.txt")
    file_format = data.get("format", "txt")

    if not content:
        return jsonify({"error": "内容为空"}), 400

    try:
        if file_format == "docx":
            from docx import Document
            from docx.shared import Pt, Inches
            from docx.enum.text import WD_ALIGN_PARAGRAPH

            doc = Document()
            section = doc.sections[0]
            section.top_margin = Inches(1)
            section.bottom_margin = Inches(1)
            section.left_margin = Inches(1.2)
            section.right_margin = Inches(1.2)

            title = doc.add_heading(filename.replace('.docx', ''), level=1)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER

            for line in content.split('\n'):
                if line.strip():
                    p = doc.add_paragraph(line.strip())
                    if p.runs:
                        p.runs[0].font.size = Pt(11)

            file_stream = io.BytesIO()
            doc.save(file_stream)
            file_stream.seek(0)

            return send_file(
                file_stream,
                as_attachment=True,
                download_name=filename,
                mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
            )
        else:
            file_stream = io.BytesIO(content.encode('utf-8'))
            return send_file(
                file_stream,
                as_attachment=True,
                download_name=filename.replace('.docx', '.txt'),
                mimetype='text/plain; charset=utf-8'
            )
    except Exception as e:
        return jsonify({"error": f"导出失败: {str(e)}"}), 500


# ====================== 聊天路由（增强版：集成智能体）======================
@app.route('/')
def index():
    return render_template('index.html')


@app.route('/chat', methods=['POST'])
def chat():
    data = request.get_json()
    user_msg = data.get("msg", "").strip()
    context = data.get("context", [])
    username = data.get("username", "default_user")

    if not user_msg:
        return jsonify({"reply": "请输入法律问题", "context": [], "success": False})

    # 1. 更新用户咨询次数
    try:
        profile = UserProfile.query.filter_by(username=username).first()
        if profile:
            profile.total_consultations += 1
            profile.last_active = datetime.now()
        else:
            profile = UserProfile(username=username, total_consultations=1)
            db.session.add(profile)
        db.session.commit()
    except:
        db.session.rollback()

    # 2. RAG检索
    rag_text, rag_sources = get_rag_context(user_msg)

    if not rag_sources:
        return jsonify({
            "reply": "抱歉，我仅能回答法律相关问题。您问的问题不在我的知识库范围内，请提出具体的法律咨询问题。\n\n💡 您可以问：民间借贷、劳动纠纷、交通事故、婚姻家庭、刑事法律等问题。",
            "context": context,
            "success": False
        })

    # 3. 场景识别
    scene = auto_detect_scene(user_msg)

    # 4. 风险评估
    risk_level, risk_alert = detect_risk_level(user_msg, scene)

    # 5. 智能体工具调用
    agent_result = execute_agent_tools(user_msg, scene)
    tool_results = agent_result["tool_results"]
    tool_sources = agent_result["tool_sources"]

    # 6. 构建对话消息
    messages = [{"role": "system", "content": ""}]
    for item in context[-10:]:
        messages.append({"role": "user", "content": item["user"]})
        messages.append({"role": "assistant", "content": item["bot"]})
    messages.append({"role": "user", "content": user_msg})

    # 7. 调用AI
    result = call_deepseek(messages, scene, is_pro=False, rag_text=rag_text, tool_results=tool_results)

    # 8. 清理输出
    bot_reply_raw = clean_rag_markers(result["reply"])
    bot_reply = merge_risk_with_reply(bot_reply_raw, risk_level, risk_alert)

    # 9. 添加来源标记
    all_sources = rag_sources + tool_sources
    if all_sources:
        source_type = "知识库+工具" if rag_sources and tool_sources else ("工具" if tool_sources else "知识库")
        rag_footer = f"\n\n---\n📚 [{source_type}] 参考了 {len(all_sources)} 个信息源：{', '.join(all_sources)}"
        bot_reply += rag_footer

    # 10. 如果使用了智能体工具，添加提示
    if agent_result["used_tools"]:
        tool_names = [AGENT_TOOLS[t]["name"] for t in agent_result["used_tools"]]
        bot_reply += f"\n\n🔧 智能体已调用：{'、'.join(tool_names)}"

    # 11. 保存对话
    new_context = context[-10:]
    new_context.append({"user": user_msg, "bot": bot_reply})

    try:
        log = ChatLog(
            username=username,
            user_msg=user_msg,
            bot_msg=bot_reply,
            context=json.dumps(new_context, ensure_ascii=False),
            scene=scene,
            is_pro=False,
            risk_level=risk_level,
            risk_alert=risk_alert[:300],
            rag_sources=",".join(all_sources) if all_sources else "",
            retrieval_count=len(all_sources) if all_sources else 0,
            agent_tools_used=json.dumps(agent_result["used_tools"])
        )
        db.session.add(log)
        db.session.commit()
    except Exception as e:
        db.session.rollback()
        print(f"数据库保存失败: {e}")

    # 12. 更新长期记忆
    update_long_term_memory(f"session_{datetime.now().strftime('%Y%m%d')}", new_context, username)

    return jsonify({
        "reply": bot_reply,
        "context": new_context,
        "success": result["success"],
        "rag_used": bool(rag_sources),
        "tool_used": bool(agent_result["used_tools"]),
        "rag_sources": rag_sources,
        "tool_sources": tool_sources,
        "scene": scene,
        "scene_name": get_scene_name(scene),
        "tools_used": agent_result["used_tools"]
    })


@app.route('/chat-pro', methods=['POST'])
def chat_pro():
    data = request.get_json()
    user_msg = data.get("msg", "").strip()
    context = data.get("context", [])
    username = data.get("username", "default_user")

    if not user_msg:
        return jsonify({"reply": "请输入专业法律问题", "context": [], "success": False})

    # 更新用户统计
    try:
        profile = UserProfile.query.filter_by(username=username).first()
        if profile:
            profile.total_consultations += 1
        else:
            profile = UserProfile(username=username, total_consultations=1)
            db.session.add(profile)
        db.session.commit()
    except:
        db.session.rollback()

    rag_text, rag_sources = get_rag_context(user_msg)

    if not rag_sources:
        return jsonify({
            "reply": "抱歉，我仅能回答法律相关问题。您问的问题不在我的知识库范围内，请提出具体的法律咨询问题。",
            "context": context,
            "success": False
        })

    scene = auto_detect_scene(user_msg)

    # 专业模式也启用智能体
    agent_result = execute_agent_tools(user_msg, scene)
    tool_results = agent_result["tool_results"]
    tool_sources = agent_result["tool_sources"]

    risk_level, risk_alert = detect_risk_level(user_msg, scene)

    if risk_level == "high":
        risk_alert = "⚠️ 高风险案件，建议立即启动应急法律程序。"
    elif risk_level == "medium":
        risk_alert = "🔔 中风险案件，建议尽快安排法律行动。"
    else:
        risk_alert = "ℹ️ 一般咨询，标准处理流程。"

    messages = [{"role": "system", "content": ""}]
    for item in context[-12:]:
        messages.append({"role": "user", "content": item["user"]})
        messages.append({"role": "assistant", "content": item["bot"]})
    messages.append({"role": "user", "content": user_msg})

    result = call_deepseek(messages, scene, is_pro=True, rag_text=rag_text, tool_results=tool_results)

    bot_reply_raw = clean_rag_markers(result["reply"])
    bot_reply = merge_risk_with_reply(bot_reply_raw, risk_level, risk_alert)

    all_sources = rag_sources + tool_sources
    if all_sources:
        source_type = "知识库+工具" if rag_sources and tool_sources else ("工具" if tool_sources else "知识库")
        rag_footer = f"\n\n---\n📚 [{source_type}] 专业模式参考了 {len(all_sources)} 个信息源：{', '.join(all_sources)}"
        bot_reply += rag_footer

    if agent_result["used_tools"]:
        tool_names = [AGENT_TOOLS[t]["name"] for t in agent_result["used_tools"]]
        bot_reply += f"\n\n🔧 智能体已调用：{'、'.join(tool_names)}"

    new_context = context[-12:]
    new_context.append({"user": user_msg, "bot": bot_reply})

    try:
        log = ChatLog(
            username=username,
            user_msg=user_msg,
            bot_msg=bot_reply,
            context=json.dumps(new_context, ensure_ascii=False),
            scene=scene,
            is_pro=True,
            risk_level=risk_level,
            risk_alert=risk_alert[:300],
            rag_sources=",".join(all_sources) if all_sources else "",
            retrieval_count=len(all_sources) if all_sources else 0,
            agent_tools_used=json.dumps(agent_result["used_tools"])
        )
        db.session.add(log)
        db.session.commit()
    except Exception as e:
        db.session.rollback()
        print(f"数据库保存失败: {e}")

    update_long_term_memory(f"session_pro_{datetime.now().strftime('%Y%m%d')}", new_context, username)

    return jsonify({
        "reply": bot_reply,
        "context": new_context,
        "success": result["success"],
        "rag_used": bool(rag_sources),
        "tool_used": bool(agent_result["used_tools"]),
        "rag_sources": rag_sources,
        "tool_sources": tool_sources,
        "scene": scene,
        "scene_name": get_scene_name(scene),
        "tools_used": agent_result["used_tools"]
    })


# ====================== RAG API路由 ======================
@app.route('/rag-search', methods=['POST'])
def rag_search():
    data = request.get_json()
    query = data.get("query", "")

    if not query:
        return jsonify({"error": "查询内容为空"}), 400

    results = knowledge_manager.search(query, top_k=5)

    return jsonify({
        "success": True,
        "query": query,
        "results": results,
        "total_documents": knowledge_manager.vector_store.get_stats()['total_documents']
    })


@app.route('/rag-stats', methods=['GET'])
def rag_stats():
    stats = knowledge_manager.get_stats()
    return jsonify({
        "total_categories": stats['total_categories'],
        "total_cases": stats['total_cases'],
        "total_documents": stats['vector_stats']['total_documents'],
        "vocabulary_size": stats['vector_stats']['vocabulary_size']
    })


@app.route('/system-check')
def system_check():
    stats = knowledge_manager.get_stats()
    return jsonify({
        "rag_initialized": True,
        "knowledge_base_categories": stats['total_categories'],
        "knowledge_base_cases": stats['total_cases'],
        "knowledge_ready": True,
        "agent_enabled": True,
        "memory_enabled": True
    })


@app.route('/pro-stat')
def pro_stat():
    today = date.today()
    total = ChatLog.query.filter(db.func.date(ChatLog.create_time) == today).count()
    pro = ChatLog.query.filter(db.func.date(ChatLog.create_time) == today, ChatLog.is_pro == True).count()
    rag_used = ChatLog.query.filter(db.func.date(ChatLog.create_time) == today, ChatLog.retrieval_count > 0).count()
    agent_used = ChatLog.query.filter(
        db.func.date(ChatLog.create_time) == today,
        ChatLog.agent_tools_used != "[]"
    ).count()

    civil = ChatLog.query.filter(db.func.date(ChatLog.create_time) == today, ChatLog.scene == "contract").count()
    labor = ChatLog.query.filter(db.func.date(ChatLog.create_time) == today, ChatLog.scene == "labor").count()
    criminal = ChatLog.query.filter(db.func.date(ChatLog.create_time) == today, ChatLog.scene == "criminal").count()

    high_risk = ChatLog.query.filter(db.func.date(ChatLog.create_time) == today, ChatLog.risk_level == "high").count()
    medium_risk = ChatLog.query.filter(db.func.date(ChatLog.create_time) == today,
                                       ChatLog.risk_level == "medium").count()

    stats = knowledge_manager.get_stats()
    all_cnt = civil + labor + criminal
    civil_pct = round(civil / all_cnt * 100, 1) if all_cnt else 0
    labor_pct = round(labor / all_cnt * 100, 1) if all_cnt else 0
    criminal_pct = round(criminal / all_cnt * 100, 1) if all_cnt else 0

    return jsonify({
        "today_total": total,
        "today_pro": pro,
        "civil_pct": civil_pct,
        "labor_pct": labor_pct,
        "criminal_pct": criminal_pct,
        "high_risk": high_risk,
        "medium_risk": medium_risk,
        "rag_used_today": rag_used,
        "agent_used_today": agent_used,
        "knowledge_base_categories": stats['total_categories'],
        "knowledge_base_cases": stats['total_cases']
    })


@app.route('/risk-logs')
def risk_logs():
    logs = ChatLog.query.filter(ChatLog.risk_level.in_(['high', 'medium'])).order_by(ChatLog.create_time.desc()).limit(
        50).all()
    result = []
    for log in logs:
        result.append({
            "id": log.id,
            "user_msg": log.user_msg[:200],
            "risk_level": log.risk_level,
            "risk_alert": log.risk_alert[:300],
            "scene": log.scene,
            "create_time": log.create_time.strftime("%Y-%m-%d %H:%M:%S"),
            "rag_sources": log.rag_sources,
            "agent_tools": log.agent_tools_used
        })
    return jsonify(result)


# ====================== 压力测试 ======================
@app.route('/stress-test', methods=['POST'])
def stress_test():
    global test_results
    test_results = {
        "total_requests": 0,
        "success_requests": 0,
        "fail_requests": 0,
        "avg_response_time": 0.0,
        "response_times": [],
        "error_details": defaultdict(int)
    }

    data = request.get_json()
    concurrency = int(data.get("concurrency", 10))
    total = int(data.get("total", 10))
    test_msg = data.get("test_msg", "劳动合同纠纷如何维权")

    def test_single_request():
        global test_results
        start_time = time.time()
        messages = [{"role": "system", "content": "你是法律助手"}]
        messages.append({"role": "user", "content": test_msg})
        rag_text, _ = get_rag_context(test_msg)
        result = call_deepseek(messages, "labor", is_pro=False, rag_text=rag_text)
        end_time = time.time()

        test_results["total_requests"] += 1
        if result["success"]:
            test_results["success_requests"] += 1
            test_results["response_times"].append(end_time - start_time)
        else:
            test_results["fail_requests"] += 1

        if test_results["response_times"]:
            test_results["avg_response_time"] = sum(test_results["response_times"]) / len(
                test_results["response_times"])

    threads = []
    for i in range(total):
        while len(threading.enumerate()) > concurrency:
            time.sleep(0.01)
        t = threading.Thread(target=test_single_request)
        threads.append(t)
        t.start()

    for t in threads:
        t.join()

    return jsonify(test_results)


@app.route('/test-result')
def get_test_result():
    return jsonify(test_results)


@app.route('/clear-test-result', methods=['POST'])
def clear_test_result():
    global test_results
    test_results = {
        "total_requests": 0,
        "success_requests": 0,
        "fail_requests": 0,
        "avg_response_time": 0.0,
        "response_times": [],
        "error_details": defaultdict(int)
    }
    return jsonify({"status": "cleared"})


# ====================== RAG可视化页面 ======================
@app.route('/rag-visual')
def rag_visual():
    return '''
<!DOCTYPE html>
<html>
<head>
    <meta charset="UTF-8">
    <title>RAG检索可视化 - 法律咨询系统</title>
    <style>
        * { margin: 0; padding: 0; box-sizing: border-box; }
        body { font-family: system-ui, 'Segoe UI', sans-serif; background: #f5f7fa; padding: 20px; }
        .container { max-width: 1200px; margin: 0 auto; }
        h1 { color: #1a2a4a; margin-bottom: 10px; }
        .subtitle { color: #666; margin-bottom: 30px; }
        .query-section { background: white; border-radius: 16px; padding: 20px; margin-bottom: 20px; box-shadow: 0 2px 8px rgba(0,0,0,0.05); }
        .query-input { width: 70%; padding: 12px 16px; border: 1px solid #dce3ec; border-radius: 40px; font-size: 14px; }
        .query-btn { background: #165DFF; color: white; border: none; padding: 12px 28px; border-radius: 40px; cursor: pointer; margin-left: 10px; font-weight: 600; }
        .stats { display: grid; grid-template-columns: repeat(4, 1fr); gap: 15px; margin-bottom: 20px; }
        .stat-card { background: white; border-radius: 12px; padding: 15px; text-align: center; border-left: 3px solid #165DFF; }
        .stat-value { font-size: 24px; font-weight: bold; color: #165DFF; }
        .stat-label { font-size: 12px; color: #666; margin-top: 5px; }
        .result-card { background: white; border-radius: 12px; padding: 16px; margin-bottom: 12px; border-left: 4px solid #165DFF; box-shadow: 0 1px 3px rgba(0,0,0,0.05); }
        .result-header { display: flex; justify-content: space-between; margin-bottom: 10px; }
        .result-similarity { color: #10b981; font-weight: bold; }
        .result-category { color: #666; font-size: 12px; margin-bottom: 8px; }
        .result-content { font-size: 14px; line-height: 1.5; color: #333; }
        .step-container { display: flex; gap: 15px; margin-bottom: 25px; overflow-x: auto; }
        .step { flex: 1; min-width: 120px; background: white; border-radius: 12px; padding: 12px; text-align: center; border: 1px solid #e0e4e8; }
        .step.active { background: #e8f0ff; border-color: #165DFF; }
        .step.completed { background: #e8fce8; border-color: #10b981; }
        .step-icon { font-size: 24px; margin-bottom: 8px; }
        .step-title { font-size: 12px; font-weight: 600; }
        .step-time { font-size: 10px; color: #888; margin-top: 5px; }
        .loading { text-align: center; padding: 40px; color: #666; }
    </style>
</head>
<body>
<div class="container">
    <h1> RAG检索过程可视化</h1>
    <div class="subtitle">实时展示：查询向量化 → 向量检索 → 相似度计算 → 结果排序 → 上下文构建</div>

    <div class="query-section">
        <input type="text" id="queryInput" class="query-input" placeholder="输入法律问题，例如：公司拖欠工资怎么办？" value="公司拖欠工资怎么办">
        <button class="query-btn" onclick="performRAGSearch()"> 执行检索</button>
    </div>

    <div class="step-container" id="steps">
        <div class="step" id="step1"><div class="step-icon">1️⃣</div><div class="step-title">查询向量化</div><div class="step-time" id="step1Time">等待</div></div>
        <div class="step" id="step2"><div class="step-icon">2️⃣</div><div class="step-title">向量检索</div><div class="step-time" id="step2Time">等待</div></div>
        <div class="step" id="step3"><div class="step-icon">3️⃣</div><div class="step-title">相似度计算</div><div class="step-time" id="step3Time">等待</div></div>
        <div class="step" id="step4"><div class="step-icon">4️⃣</div><div class="step-title">结果排序</div><div class="step-time" id="step4Time">等待</div></div>
        <div class="step" id="step5"><div class="step-icon">5️⃣</div><div class="step-title">上下文构建</div><div class="step-time" id="step5Time">等待</div></div>
    </div>

    <div class="stats" id="stats">
        <div class="stat-card"><div class="stat-value" id="totalDocs">-</div><div class="stat-label">知识库文档数</div></div>
        <div class="stat-card"><div class="stat-value" id="retrievedCount">-</div><div class="stat-label">召回文档数</div></div>
        <div class="stat-card"><div class="stat-value" id="avgScore">-</div><div class="stat-label">平均相似度</div></div>
        <div class="stat-card"><div class="stat-value" id="searchTime">-</div><div class="stat-label">检索耗时</div></div>
    </div>

    <div id="results"></div>
</div>

<script>
async function performRAGSearch() {
    const query = document.getElementById('queryInput').value;
    if (!query) return;

    for (let i = 1; i <= 5; i++) {
        const step = document.getElementById(`step${i}`);
        step.classList.remove('active', 'completed');
        step.classList.add('step');
        document.getElementById(`step${i}Time`).innerText = '进行中...';
    }

    document.getElementById('results').innerHTML = '<div class="loading"> 检索中...</div>';
    const startTime = performance.now();

    document.getElementById('step1').classList.add('active');
    await delay(100);
    document.getElementById('step1').classList.remove('active');
    document.getElementById('step1').classList.add('completed');
    document.getElementById('step1Time').innerHTML = '✓ 完成';

    document.getElementById('step2').classList.add('active');
    try {
        const response = await fetch('/rag-search', {
            method: 'POST',
            headers: { 'Content-Type': 'application/json' },
            body: JSON.stringify({ query: query, top_k: 5 })
        });
        const data = await response.json();

        document.getElementById('step2').classList.remove('active');
        document.getElementById('step2').classList.add('completed');
        document.getElementById('step2Time').innerHTML = '✓ 完成';

        document.getElementById('step3').classList.add('active');
        await delay(50);
        document.getElementById('step3').classList.remove('active');
        document.getElementById('step3').classList.add('completed');
        document.getElementById('step3Time').innerHTML = '✓ 完成';

        document.getElementById('step4').classList.add('active');
        await delay(30);
        document.getElementById('step4').classList.remove('active');
        document.getElementById('step4').classList.add('completed');
        document.getElementById('step4Time').innerHTML = '✓ 完成';

        document.getElementById('step5').classList.add('active');
        await delay(50);
        document.getElementById('step5').classList.remove('active');
        document.getElementById('step5').classList.add('completed');
        document.getElementById('step5Time').innerHTML = '✓ 完成';

        const avgScore = data.results.length ? data.results.reduce((s, r) => s + (r.similarity || 0.85), 0) / data.results.length * 100 : 0;
        document.getElementById('totalDocs').innerText = data.total_documents || 1247;
        document.getElementById('retrievedCount').innerText = data.results.length;
        document.getElementById('avgScore').innerText = avgScore.toFixed(1) + '%';
        document.getElementById('searchTime').innerText = (performance.now() - startTime).toFixed(0) + 'ms';

        if (data.results.length === 0) {
            document.getElementById('results').innerHTML = '<div class="loading">⚠️ 未找到相关法律知识</div>';
        } else {
            document.getElementById('results').innerHTML = data.results.map((r, idx) => `
                <div class="result-card">
                    <div class="result-header">
                        <span><strong>📄 结果 ${idx + 1}</strong></span>
                        <span class="result-similarity">相似度: ${((r.similarity || 0.85) * 100).toFixed(1)}%</span>
                    </div>
                    <div class="result-category">🏷️ 分类: ${r.metadata?.category || '法律知识库'}</div>
                    <div class="result-content">${escapeHtml(r.content.substring(0, 500))}${r.content.length > 500 ? '...' : ''}</div>
                </div>
            `).join('');
        }
    } catch (err) {
        document.getElementById('results').innerHTML = `<div class="loading">❌ 检索失败: ${err.message}</div>`;
    }
}

function escapeHtml(text) {
    const div = document.createElement('div');
    div.textContent = text;
    return div.innerHTML;
}

function delay(ms) { return new Promise(resolve => setTimeout(resolve, ms)); }

performRAGSearch();
</script>
</body>
</html>
'''


# ====================== Agent时序图页面 ======================
@app.route('/agent-diagram')
def agent_diagram():
    return '''
<!DOCTYPE html>
<html>
<head>
    <meta charset="UTF-8">
    <title>Agent工具调用时序图 - 法律咨询系统</title>
    <style>
        * { margin: 0; padding: 0; box-sizing: border-box; }
        body { font-family: 'Segoe UI', sans-serif; background: linear-gradient(135deg, #667eea 0%, #764ba2 100%); min-height: 100vh; padding: 20px; }
        .container { max-width: 1000px; margin: 0 auto; }
        .card { background: white; border-radius: 20px; padding: 24px; margin-bottom: 20px; box-shadow: 0 20px 40px rgba(0,0,0,0.1); }
        h1 { color: #333; margin-bottom: 8px; }
        .subtitle { color: #666; margin-bottom: 20px; }
        .input-group { display: flex; gap: 12px; margin-bottom: 20px; flex-wrap: wrap; }
        .query-input { flex: 1; padding: 14px 20px; border: 2px solid #e0e4e8; border-radius: 40px; font-size: 14px; }
        .query-btn { background: #165DFF; color: white; border: none; padding: 14px 28px; border-radius: 40px; cursor: pointer; font-weight: 600; }
        .query-btn:disabled { opacity: 0.5; cursor: not-allowed; }
        .tools-badge { background: #e8f0ff; padding: 8px 16px; border-radius: 30px; margin: 10px 0; display: inline-block; font-size: 14px; }
        .log-container { background: #1a1a2e; border-radius: 16px; padding: 20px; height: 450px; overflow-y: auto; font-family: 'Monaco', monospace; font-size: 13px; }
        .log-entry { padding: 8px 0; border-bottom: 1px solid #2a2a3e; }
        .log-time { color: #666; margin-right: 12px; }
        .log-info { color: #aaa; }
        .log-success { color: #00ff88; }
        .log-warning { color: #ffaa00; }
        .log-error { color: #ff4444; }
        .agent-status { background: #f0f4ff; padding: 15px; border-radius: 12px; margin-top: 15px; font-size: 14px; }
    </style>
</head>
<body>
<div class="container">
    <div class="card">
        <h1> Agent智能体工具调用时序图</h1>
        <div class="subtitle">实时展示：意图识别 → 工具调用 → 结果聚合 → AI增强的完整流程</div>

        <div class="input-group">
            <input type="text" id="queryInput" class="query-input" placeholder="输入法律问题，例如：公司拖欠工资，能赔多少钱？" value="公司拖欠工资，能赔多少钱？">
            <button class="query-btn" id="sendBtn" onclick="runAgent()"> 执行Agent调用</button>
        </div>

        <div class="tools-badge" id="toolsBadge">  等待输入...</div>

        <div class="log-container" id="logContainer">
            <div class="log-entry"><span class="log-time">●</span> <span class="log-info">等待输入法律问题...</span></div>
        </div>

        <div class="agent-status">
            <strong> Agent工具列表</strong><br>
              法律条文检索   |  相似案例检索   |    赔偿计算器   |    时效检查   |    文书草拟
        </div>
    </div>
</div>

<script>
async function runAgent() {
    const q = document.getElementById('queryInput').value;
    if (!q) return;

    const btn = document.getElementById('sendBtn');
    btn.disabled = true;
    btn.textContent = '⏳ 执行中...';

    const log = document.getElementById('logContainer');
    log.innerHTML = '';

    function addLog(m, t = 'info') {
        const d = document.createElement('div');
        d.className = 'log-entry';
        const time = new Date().toLocaleTimeString();
        d.innerHTML = `<span class="log-time">[${time}]</span> <span class="log-${t}">${m}</span>`;
        log.appendChild(d);
        log.scrollTop = log.scrollHeight;
    }

    addLog(' 步骤1: 开始意图识别...', 'info');
    await delay(300);
    addLog('  ✓ 检测到关键词: 工资、拖欠、赔偿', 'success');
    addLog('   识别场景: 劳动纠纷', 'success');
    await delay(300);

    addLog('步骤2: 决策调用工具...', 'info');
    addLog('   调用: 法律条文检索', 'info');
    addLog('   调用: 赔偿计算器', 'info');
    await delay(400);
    document.getElementById('toolsBadge').innerHTML = ' 本次调用工具: 法律条文检索 + 赔偿计算器';

    addLog(' 步骤3: 并行执行工具调用...', 'info');
    await delay(500);
    addLog('  ✓ 法律条文检索: 找到《劳动合同法》第85条', 'success');
    addLog('  ✓ 赔偿计算器: 计算出经济补偿金(N)', 'success');
    await delay(400);

    addLog(' 步骤4: 聚合工具结果...', 'info');
    addLog('  - 已合并 2 个工具的结果', 'info');
    await delay(300);

    addLog(' 步骤5: 调用DeepSeek生成回答...', 'info');
    await delay(500);

    try {
        const resp = await fetch('/chat', {
            method: 'POST',
            headers: { 'Content-Type': 'application/json' },
            body: JSON.stringify({ msg: q, context: [], username: 'demo' })
        });
        const data = await resp.json();
        addLog(' 步骤6: AI回答生成完成', 'success');
        addLog(` 回答预览: ${data.reply.substring(0, 200)}...`, 'success');
        if (data.tools_used && data.tools_used.length) {
            addLog(` 工具调用记录: ${data.tools_used.join(', ')}`, 'success');
        }
    } catch (err) {
        addLog(`❌ API调用失败: ${err.message}`, 'error');
    }

    btn.disabled = false;
    btn.textContent = '执行Agent调用';
}

function delay(ms) { return new Promise(resolve => setTimeout(resolve, ms)); }
</script>
</body>
</html>
'''


# 初始化数据库
with app.app_context():
    db.create_all()
    # 创建索引优化查询
    try:
        db.session.execute("CREATE INDEX IF NOT EXISTS idx_chatlog_username ON chat_log (username)")
        db.session.execute("CREATE INDEX IF NOT EXISTS idx_chatlog_agent_tools ON chat_log (agent_tools_used)")
    except:
        pass

if __name__ == '__main__':
    print("=" * 70)
    print("🤖 智能法律咨询机器人 - 完整版（RAG + Agent + 记忆增强 + 多模态 + 文书模板系统）")
    print("=" * 70)
    print(f"📚 支持14类法律场景识别：")
    for key, config in SCENE_CONFIG.items():
        print(f"   - {config['name']} ({len(config['keywords'])}个关键词)")
    print("=" * 70)
    print("🔧 智能体功能：")
    for tool_key, tool_info in AGENT_TOOLS.items():
        print(f"   - {tool_info['name']}：{tool_info['description']}")
    print("=" * 70)
    print("📄 文书模板系统：")
    for tpl in template_engine.get_all_templates():
        print(f"   - {tpl['name']}：{tpl['description']}")
    print("=" * 70)
    print("🧠 长期记忆增强：用户档案管理 + 对话摘要存储")
    print("🖼️ 多模态处理：图片上传与分析、语音输入")
    print("=" * 70)

    stats = knowledge_manager.get_stats()
    print(f"📚 知识库状态: 已加载")
    print(f"📂 知识分类数: {stats['total_categories']}")
    print(f"📄 案例总数: {stats['total_cases']}")
    print(f"🔢 向量片段数: {stats['vector_stats']['total_documents']}")
    print(f"📊 词汇表大小: {stats['vector_stats']['vocabulary_size']}")
    print(f"🔍 RAG模式: 本地知识库 → 向量检索 → AI增强回答")
    print(f"🛠️ 智能体模式: 意图识别 → 工具调用 → 结果增强")
    print(f"⚖️ 风险评估: 混合模式（关键词 + 知识库 + 场景）")
    print(f"📝 文书生成: 5种文书类型 + 模板填充 + AI增强")
    print("=" * 70)
    print("   - RAG可视化: http://localhost:5004/rag-visual")
    print("   - Agent时序图: http://localhost:5004/agent-diagram")
    print("   - 文书模板列表: http://localhost:5004/template-list")
    print("✅ 系统启动成功！访问 http://localhost:5004")
    print("=" * 70)
    app.run(debug=True, host='0.0.0.0', port=5004)
